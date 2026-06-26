import urllib.request
import base64
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def download_mermaid_ink(mermaid_str, filename):
    # mermaid.ink uses base64 string directly
    encoded = base64.b64encode(mermaid_str.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}'
    
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req, timeout=15) as response, open(filename, 'wb') as out_file:
            out_file.write(response.read())
        print(f'Downloaded {filename}')
        return True
    except Exception as e:
        print(f'Failed to download {filename}: {e}')
        return False

# Define diagrams
architecture_code = """
graph TD
    subgraph Mobile Client
        UI[Flutter App]
    end
    subgraph Backend Server
        API[Flask REST API]
        DB[(SQLite Database)]
        API --> DB
    end
    subgraph Intelligence Layer
        RN[ResNet50 Model]
        KNN[KNN Soil Engine]
        LSTM[LSTM Forecaster]
        LLM[Groq LLaMA 3.3]
    end
    UI <-->|JSON / HTTPS| API
    API <-->|Inference| RN
    API <-->|Analysis| KNN
    API <-->|Prediction| LSTM
    API <-->|API Calls| LLM
"""

er_code = """
erDiagram
    USER ||--o{ LISTING : "creates"
    USER ||--o{ PURCHASE : "makes"
    USER ||--o{ MESSAGE : "sends"
    USER ||--o{ TRANSACTION : "has"
    USER ||--o{ CALENDARTASK : "manages"
    USER ||--o{ ECOMMERCEORDER : "places"
    LISTING ||--o{ PURCHASE : "contains"
    PURCHASE ||--o{ MESSAGE : "tracks"
"""

dfd_code = """
graph LR
    F((Farmer)) -->|Images, Soil Data| System[KrishiMitra System]
    B((Buyer)) -->|Search, Bids| System
    System -->|Recommendations| F
    System -->|Listings| B
    System <-->|Prices| Agmarknet[Agmarknet API]
    System <-->|Advisory| Groq[Groq API]
"""

flowchart_code = """
graph TD
    A([Start]) --> B[Capture Leaf Image]
    B --> C[Compress & Encode Base64]
    C --> D[Transmit via HTTPS POST]
    D --> E{Backend Flask API}
    E --> F[Decode Image Tensor]
    F --> G[ResNet50 Inference]
    G --> H[Extract Feature Maps]
    H --> I[Apply Grad-CAM]
    I --> J[Overlay JET Colormap]
    J --> K[Return JSON + Image URL]
    K --> L[Render Heatmap]
    L --> M([End])
"""

diagrams = [
    (architecture_code, "diagram_arch.png", "System Architecture Design"),
    (er_code, "diagram_er.png", "Entity-Relationship (ER) Diagram"),
    (dfd_code, "diagram_dfd.png", "Data Flow Diagram (Level 0 Context Diagram)"),
    (flowchart_code, "diagram_flow.png", "Flowchart: Disease Detection Pipeline")
]

# Download images
downloaded = []
for code, filename, caption in diagrams:
    if download_mermaid_ink(code, filename):
        downloaded.append((filename, caption))

# Append to document
doc_path = "KrishiMitra_Final_Report_Massive_With_Figures.docx"
print(f"Opening {doc_path}...")
try:
    doc = Document(doc_path)
except Exception as e:
    print(f"Error opening doc: {e}")
    exit(1)

doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Appendix D: System Diagrams and Architecture")
r.bold = True
r.font.size = Pt(14)
r.font.name = "Times New Roman"

p2 = doc.add_paragraph()
r2 = p2.add_run("The following architectural diagrams, Entity-Relationship (ER) models, Data Flow Diagrams (DFDs), and algorithmic flowcharts detail the internal structural logic and data pathways of the KrishiMitra platform.")
r2.font.size = Pt(12)
r2.font.name = "Times New Roman"
p2.paragraph_format.line_spacing = Pt(21)

fig_idx = 1
for filename, caption in downloaded:
    # Add Image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(filename, width=Inches(5.5))
    
    # Add Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run(f"Figure D.{fig_idx}: {caption}")
    r_cap.font.size = Pt(12)
    r_cap.font.name = "Times New Roman"
    r_cap.bold = True
    fig_idx += 1
    
    # Empty line
    doc.add_paragraph()

out_name = "KrishiMitra_Final_Report_Massive_FINAL_COMPLETE.docx"
doc.save(out_name)
print(f"Successfully generated diagrams and saved as {out_name}!")
