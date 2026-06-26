import copy
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_borders(doc):
    for sec in doc.sections:
        sectPr = sec._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'double')
            border.set(qn('w:sz'), '12') # Size of border (1/8 pt)
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
            
        sectPr.append(pgBorders)

def add_headers(doc):
    for sec in doc.sections:
        header = sec.header
        # Clear existing
        for p in header.paragraphs:
            p.text = ""
        
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        r = p.add_run("AI-Powered Smart Agriculture Ecosystem for Small and Marginal Farmers.")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a bottom border to the header paragraph
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

# 1. Merge Documents
print("Merging documents...")
doc_front = Document("KrishiMitra_Final_Report.docx")
doc_massive = Document("KrishiMitra_Final_Report_Massive.docx")

# Remove Chapter 1 and onwards from front document
start_deleting = False
for element in list(doc_front.element.body):
    if element.tag.endswith('p') and "Chapter 1" in element.text:
        start_deleting = True
    if start_deleting:
        doc_front.element.body.remove(element)

# Append massive document
for element in doc_massive.element.body:
    # Do not copy section properties at the end of the body
    if not element.tag.endswith('sectPr'):
        doc_front.element.body.append(copy.deepcopy(element))

# 2. Add Headers and Borders
print("Adding headers and borders...")
add_headers(doc_front)
add_page_borders(doc_front)

# Save intermediate
temp_name = "temp_master.docx"
doc_front.save(temp_name)

# 3. Insert Diagrams
print("Inserting diagrams...")
doc = Document(temp_name)

diagrams = [
    ("real_arch.png", "System Architecture Design", "1.4 Proposed Approach", 1),
    ("real_dfd.png", "Data Flow Diagram (Level 0)", "5.3.2 Data Flow Diagrams (DFD)", 2),
    ("real_er.png", "Entity-Relationship (ER) Diagram", "5.4 Database Design and Data Dictionary", 3),
    ("real_flow.png", "Flowchart: Disease Detection Pipeline", "6.2.1 Deep Learning Vision Module", 4)
]

for filename, caption, target_text, fig_num in diagrams:
    if not os.path.exists(filename):
        print(f"Warning: {filename} not found, skipping insertion.")
        continue
        
    for i, p in enumerate(doc.paragraphs):
        if target_text in p.text:
            insert_idx = min(i + 2, len(doc.paragraphs) - 1)
            target_p = doc.paragraphs[insert_idx]
            
            new_p = target_p.insert_paragraph_before()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_p.add_run()
            run.add_picture(filename, width=Inches(5.5))
            
            cap_p = target_p.insert_paragraph_before()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_r = cap_p.add_run(f"Figure {fig_num}: {caption}")
            cap_r.font.size = Pt(12)
            cap_r.font.name = "Times New Roman"
            cap_r.bold = True
            
            target_p.insert_paragraph_before("")
            print(f"Inserted {filename} into {target_text}")
            break

out_name = "KrishiMitra_MASTER_REPORT.docx"
doc.save(out_name)
print(f"Successfully generated {out_name}!")
