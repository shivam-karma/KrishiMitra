from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def new_doc():
    doc = Document()
    for sec in doc.sections:
        sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
        sec.left_margin = Cm(2.54); sec.right_margin = Cm(2.54)
        sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    return doc

def para(doc, text, size=12, bold=False, align="JUSTIFY", italic=False):
    p = doc.add_paragraph()
    p.alignment = {"CENTER": WD_ALIGN_PARAGRAPH.CENTER, "LEFT": WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.line_spacing = Pt(21)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = "Times New Roman"
    return p

def chap(d,t): para(d,t,16,True,"CENTER")
def sec(d,t):  para(d,t,14,True,"LEFT")
def sub(d,t):  para(d,t,12,True,"LEFT")
def body(d,t): para(d,t,12,align="JUSTIFY")
def cap(d,t):  para(d,t,11,False,"CENTER",True)
def blank(d):  para(d,"",align="LEFT")
def br(d):     d.add_page_break()

def tbl(doc, headers, rows, caption=""):
    if caption:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption); r.bold=True; r.font.size=Pt(11); r.font.name="Times New Roman"
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style="Table Grid"
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h
        for run in c.paragraphs[0].runs: run.bold=True; run.font.size=Pt(10); run.font.name="Times New Roman"; run.font.color.rgb=RGBColor(255,255,255)
        tc=c._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:fill'),'1F3864'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]; fill='EEF4FF' if ri%2==0 else 'FFFFFF'
        for ci,val in enumerate(row):
            c=tr.cells[ci]; c.text=str(val)
            for p in c.paragraphs:
                for run in p.runs: run.font.size=Pt(10); run.font.name="Times New Roman"
            tc=c._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
            shd.set(qn('w:fill'),fill); shd.set(qn('w:val'),'clear'); tcPr.append(shd)
    blank(doc)

doc = new_doc()

chap(doc,"Chapter 5")
chap(doc,"Analysis and Design")
blank(doc)

body(doc,"Analysis and design are the two foundational pillars upon which the KrishiMitra system is constructed. The analysis phase identifies what the system must do — capturing functional and non-functional requirements, data flows, and user interactions — while the design phase defines how the system will fulfil those requirements through architectural decisions, module decomposition, communication models, and interface specifications. This chapter presents the complete system analysis and design of KrishiMitra, including requirements specification, system block diagram, data flow diagrams, use case diagrams, software stack selection, standards compliance, and the IoT-equivalent architectural decomposition adapted for a mobile AI platform [1].")
blank(doc)

sec(doc,"5.1 Requirements Specification")
body(doc,"The requirements of KrishiMitra were captured through structured analysis of the five project objectives defined in Chapter 1. Requirements are classified into hardware requirements, software requirements, data requirements, security requirements, and user interface requirements. Table 5.1 summarises the system requirements specification for KrishiMitra.")
blank(doc)

tbl(doc,
    ["Requirement Category","Specification"],
    [
        ["Purpose","A cross-platform mobile application that provides AI-powered agricultural advisory services — disease detection, soil crop recommendation, weather advisory, P2P marketplace, and AI chatbot — to Indian smallholder farmers via a Flutter mobile client and Flask REST API backend."],
        ["Behaviour","The system operates in two roles: Farmer (can list crops, detect diseases, analyse soil, view weather) and Buyer (can browse listings, purchase crops, chat with farmers). All features require OTP-authenticated login. Disease detection and soil analysis are AI-driven and return results in under 3 seconds."],
        ["System Hardware Requirements","Android device with minimum 2GB RAM, 16GB storage, 8MP rear camera (for disease detection image capture), and 4G/Wi-Fi connectivity. Backend server: Python 3.10, 4GB RAM minimum, with GPU optional for CNN inference."],
        ["System Software Requirements","Flutter SDK 3.x, Dart 3.x, Android SDK API level 21+, Python 3.10, Flask 2.x, SQLAlchemy 2.x, TensorFlow 2.x, OpenCV 4.x, NumPy, Requests library."],
        ["Data Collection Requirements","PlantVillage dataset (54,306 images, 38 disease classes) for CNN training. Custom soil dataset (2,200 samples, 22 crop classes, features: N, P, K, pH, moisture) for KNN training. Real-time weather data from OpenWeatherMap API."],
        ["Data Analysis Requirements","CNN model: image classification with confidence score and Grad-CAM heatmap. KNN model: crop recommendation with confidence %, soil health score (0-100), fertilizer advice. LSTM: price trend forecasting (future scope)."],
        ["System Management Requirements","Role-based access control (Farmer/Buyer). Session persistence via SharedPreferences. Admin can view all listings and users via backend. Automated OTP expiry after 10 minutes."],
        ["Security Requirements","OTP-based phone authentication via Fast2SMS. Passwords hashed using bcrypt. JWT or session tokens for authenticated API calls. HTTPS for production deployment. CORS restricted to Flutter app origin."],
        ["User Interface Requirements","Material Design 3 UI with dark/light theme support. Bottom navigation bar with 5 primary sections. Loading indicators for all async AI/API calls. Error messages for network failures. Confidence score displayed with colour coding (green >80%, amber 50-80%, red <50%)."],
    ],
    "Table 5.1 KrishiMitra System Requirements Specification"
)

sec(doc,"5.2 System Block Diagram")
body(doc,"Figure 5.1 illustrates the functional block diagram of the KrishiMitra system. The diagram follows the standard convention of input blocks on the left, processing blocks in the centre, and output blocks on the right. The three primary input sources are: (1) the Farmer/Buyer mobile client providing image uploads, soil parameter inputs, and text queries; (2) external APIs providing real-time weather data and OTP delivery; and (3) the AI training datasets loaded during model initialisation. The central processing block is the Flask REST API server, which routes requests to the appropriate intelligence module (CNN for disease, KNN for soil, Groq LLaMA for chatbot, OpenWeatherMap for weather). The output blocks represent the Flutter UI screens that render the AI-generated advisory content to the end user [2].")
blank(doc)
cap(doc,"Figure 5.1: Functional Block Diagram of KrishiMitra System")
blank(doc)
body(doc,"The block diagram clearly demonstrates the three-tier architecture of KrishiMitra. Tier 1 (Presentation): the Flutter mobile application renders dynamic UI based on JSON responses from the backend. Tier 2 (Application Logic): the Flask backend hosts all business logic, AI model inference, and external API integrations. Tier 3 (Data): the SQLite database (managed by SQLAlchemy ORM) provides persistent storage for user profiles, crop listings, purchase records, chat messages, and financial transactions. This separation of concerns ensures that each tier can be independently scaled, tested, and maintained [3].")
blank(doc)

sec(doc,"5.3 System Flowchart")
body(doc,"Figure 5.2 presents the system flowchart for the KrishiMitra application. The flow begins with the application launch, which checks for an existing authenticated session stored in SharedPreferences. If a session is found, the user is directed to the role-appropriate dashboard (Farmer Dashboard or Buyer Dashboard). If no session exists, the user is directed to the Login screen, where they enter their phone number to receive an OTP via Fast2SMS. Upon successful OTP verification, the user selects their role (Farmer or Buyer), and the session is persisted locally. From the dashboard, the user can navigate to any of the eight modules. Each AI module (Disease Detection, Soil Analysis) follows an asynchronous request-response pattern: the Flutter client submits data to the Flask API, receives a JSON response containing the prediction and advisory, and renders it on the result screen. The marketplace module follows a state-machine pattern: a listing progresses through states — Available → Purchased → Confirmed → Completed — driven by buyer and farmer actions [4].")
blank(doc)
cap(doc,"Figure 5.2: System Flowchart for KrishiMitra Application")
blank(doc)

sec(doc,"5.4 Technology Stack Selection")
body(doc,"The selection of each technology component in KrishiMitra was driven by objective comparative analysis against project requirements. Table 5.2 compares the evaluated mobile frontend frameworks, justifying the selection of Flutter.")
blank(doc)

tbl(doc,
    ["Feature/Criterion","Flutter (Dart)","React Native (JS)","Android Native (Java/Kotlin)","Ionic (HTML/JS)"],
    [
        ["Cross-platform (Android & iOS)","Yes — single codebase","Yes — single codebase","No — Android only","Yes — single codebase"],
        ["Performance","Near-native (compiled to ARM)","Near-native (JS bridge)","Native","Web-based (slower)"],
        ["UI Rendering","Custom Skia/Impeller engine","Native components","Native components","WebView-based"],
        ["Camera/File access","Flutter camera plugin (easy)","React Native camera (complex)","Built-in (easy)","Cordova plugin (limited)"],
        ["Community & Packages","Large, growing (pub.dev)","Very large (npm)","Mature","Smaller"],
        ["Hot reload for development","Yes","Yes","Limited","Yes"],
        ["Suitability for AI result screens","Excellent — custom widgets","Good","Good","Poor — limited animation"],
        ["Selected","✅ Yes","No","No","No"],
    ],
    "Table 5.2 Comparison of Mobile Frontend Frameworks"
)

body(doc,"Table 5.3 compares the evaluated Python backend frameworks, justifying the selection of Flask.")
blank(doc)

tbl(doc,
    ["Feature/Criterion","Flask","Django","FastAPI","Express.js (Node)"],
    [
        ["Language","Python","Python","Python","JavaScript"],
        ["Architecture","Micro-framework (lightweight)","Full-stack (heavyweight)","Micro-framework (async)","Minimal (middleware-based)"],
        ["TensorFlow integration","Easy (native Python)","Easy (native Python)","Easy (native Python)","Complex (child process)"],
        ["ORM support","SQLAlchemy (via extension)","Built-in Django ORM","SQLAlchemy (via extension)","Sequelize (separate library)"],
        ["REST API creation","Simple with Flask-RESTful","DRF (verbose)","Automatic OpenAPI docs","Express Router"],
        ["Learning curve","Low","Medium-High","Low-Medium","Low"],
        ["Production deployment","Gunicorn + Nginx","Gunicorn + Nginx","Uvicorn + Nginx","PM2 + Nginx"],
        ["Selected","✅ Yes","No","No","No"],
    ],
    "Table 5.3 Comparison of Backend Frameworks"
)

sec(doc,"5.5 Standards and Protocols")
body(doc,"KrishiMitra adheres to a comprehensive set of industry standards and communication protocols that ensure interoperability, security, and reliability of the system. Compliance with these standards is essential for a production-grade application serving sensitive agricultural and financial data [5].")
blank(doc)

tbl(doc,
    ["Standard/Protocol","Category","Application in KrishiMitra"],
    [
        ["HTTP/HTTPS (RFC 7230-7235)","Communication Protocol","All Flutter-to-Flask API communication uses HTTP/1.1; HTTPS (TLS 1.2+) required for production deployment to encrypt data in transit"],
        ["REST (Representational State Transfer)","API Architecture","All 20+ backend endpoints follow REST principles: stateless, resource-based URLs (/users, /listings, /detect-disease), standard HTTP verbs (GET, POST, PUT, DELETE)"],
        ["JSON (RFC 8259)","Data Format","All API request and response bodies are JSON-encoded; consistent schema with 'status', 'data', and 'message' fields"],
        ["JWT (JSON Web Tokens, RFC 7519)","Security","Session tokens for authenticated API requests; tokens include user_id, role, and expiry timestamp"],
        ["bcrypt (IETF RFC 4648)","Security","Password hashing for user credentials stored in SQLite database; salt rounds = 12"],
        ["OTP (RFC 6238 — TOTP)","Authentication","Time-based 6-digit OTP generated server-side with 10-minute expiry; delivered via Fast2SMS SMS gateway"],
        ["Material Design 3 (Google)","UI Standard","Flutter frontend follows Material Design 3 guidelines for typography, colour system, component library, and interaction patterns"],
        ["TensorFlow SavedModel format","AI/ML Standard","CNN disease detection model serialised in TensorFlow SavedModel format for portable deployment and version management"],
        ["ISO/IEC 27001","Information Security","Security controls applied: access control, cryptography (bcrypt, HTTPS), logging, and incident response planning"],
        ["DPDPA 2023 (India)","Data Privacy Law","User phone numbers collected with consent; stored securely; OTP not logged; no third-party data sharing; privacy policy displayed at registration"],
        ["OpenAPI 3.0 (Swagger)","API Documentation","API endpoints documented in OpenAPI 3.0 format for developer reference and Postman import"],
    ],
    "Table 5.4 Standards and Protocols Applied in KrishiMitra"
)

sec(doc,"5.6 System Architecture — Three-Tier Model")
body(doc,"Figure 5.3 illustrates the three-tier system architecture of KrishiMitra. The Presentation Tier consists of the Flutter mobile application, which includes 15 screens organised around a bottom navigation bar. The Application Logic Tier is the Flask REST API backend, which hosts all business logic and AI inference. The Data Tier consists of the SQLite database managed via SQLAlchemy ORM, containing seven tables: users, listings, purchases, messages, finance_records, calendar_tasks, and soil_results. This architecture enforces strict separation of concerns: the Flutter client is purely a presentation layer with no business logic; all intelligence resides in the Flask backend; the database is accessed exclusively through the ORM, ensuring schema integrity [6].")
blank(doc)
cap(doc,"Figure 5.3: Three-Tier System Architecture of KrishiMitra")
blank(doc)

sec(doc,"5.7 Use Case Diagram")
body(doc,"Figure 5.4 presents the Use Case Diagram for KrishiMitra. The system has three actors: Farmer, Buyer, and System (representing automated background processes). The Farmer actor can perform: Register/Login (OTP), Detect Disease, Analyse Soil, View Weather Advisory, List Crop for Sale, Chat with Buyer, Track Finances, View Farming Calendar, Consult AI Chatbot, and Access Government Schemes. The Buyer actor can perform: Register/Login (OTP), Browse Crop Listings, Purchase Crop, Chat with Farmer, and View Crop Details. The System actor performs: Generate OTP, Send SMS, Train ML Models (during setup), and Fetch Weather Data. Key include relationships exist between Login and Generate OTP, and between Detect Disease and Process CNN Inference [7].")
blank(doc)
cap(doc,"Figure 5.4: Use Case Diagram for KrishiMitra")
blank(doc)

sec(doc,"5.8 Data Flow Diagram — Context Level (Level 0)")
body(doc,"Figure 5.5 shows the Level 0 Data Flow Diagram (Context Diagram) for KrishiMitra. At the context level, KrishiMitra is represented as a single process (the System) interacting with four external entities: (1) Farmer — provides crop images, soil parameters, queries, and financial data; receives disease reports, crop recommendations, weather advisories, chatbot responses, and marketplace listings; (2) Buyer — submits search queries and purchase requests; receives listing details and purchase confirmations; (3) Fast2SMS Gateway — receives OTP delivery requests from the system; sends OTP SMS to users; (4) OpenWeatherMap API — receives location-based weather requests; returns current conditions and 5-day forecasts [8].")
blank(doc)
cap(doc,"Figure 5.5: Level 0 Data Flow Diagram (Context Diagram) for KrishiMitra")
blank(doc)

sec(doc,"5.9 Data Flow Diagram — Level 1")
body(doc,"Figure 5.6 presents the Level 1 Data Flow Diagram, which decomposes the KrishiMitra system into its seven primary processes: (P1) Authentication Manager, (P2) Disease Detection Engine, (P3) Soil Analysis Engine, (P4) Weather Advisory Module, (P5) Marketplace Manager, (P6) AI Chatbot Interface, and (P7) Farm Management Hub. Each process interacts with the central SQLite Data Store and with external entities as appropriate. The Authentication Manager (P1) handles all OTP generation, verification, and session management. The Disease Detection Engine (P2) receives image data from the Farmer, passes it to the CNN model, and returns a structured diagnostic report. The Soil Analysis Engine (P3) applies the KNN algorithm to soil parameter inputs and returns crop recommendations. The Marketplace Manager (P5) mediates all transactions between Farmers and Buyers, including listing creation, purchase requests, and chat message routing [9].")
blank(doc)
cap(doc,"Figure 5.6: Level 1 Data Flow Diagram for KrishiMitra")
blank(doc)

sec(doc,"5.10 Communication Model")
body(doc,"KrishiMitra employs the Request-Response communication model for all interactions between the Flutter mobile client and the Flask backend. In this model, the Flutter client (acting as the requester) initiates all communication by sending an HTTP request to a specific API endpoint. The Flask server (acting as the responder) processes the request synchronously and returns an HTTP response containing the result as a JSON payload. This model is appropriate for KrishiMitra because all AI inference operations (CNN disease detection, KNN soil recommendation) are request-driven — they are triggered by explicit user actions (submitting an image, entering soil parameters) rather than by continuous data streams [10].")
blank(doc)
cap(doc,"Figure 5.7: Request-Response Communication Model for KrishiMitra")
blank(doc)

body(doc,"The in-app marketplace chat feature uses a polling-based quasi-real-time communication pattern, where the Flutter client polls the /messages/{listing_id} endpoint every 5 seconds to retrieve new messages. This approach avoids the complexity of WebSocket implementation while providing acceptable message latency for agricultural negotiation use cases. A full WebSocket (publish-subscribe) model is identified as a future enhancement in Chapter 9.")
blank(doc)

sec(doc,"5.11 Database Design")
body(doc,"Table 5.5 presents the complete database schema for KrishiMitra's SQLite database, implemented using SQLAlchemy ORM. The database consists of seven tables with clearly defined primary keys, foreign key relationships, and data type constraints. The schema was designed to be normalised to 3NF (Third Normal Form) to eliminate data redundancy while maintaining query efficiency for the expected data volumes.")
blank(doc)

tbl(doc,
    ["Table","Column","Data Type","Constraint","Description"],
    [
        ["users","id","Integer","PRIMARY KEY, AUTO","Unique user identifier"],
        ["users","phone","String(15)","UNIQUE, NOT NULL","Phone number used for OTP authentication"],
        ["users","name","String(100)","NOT NULL","User's display name"],
        ["users","role","String(10)","NOT NULL","'farmer' or 'buyer'"],
        ["users","password_hash","String(255)","NOT NULL","bcrypt-hashed password"],
        ["users","created_at","DateTime","DEFAULT NOW()","Account creation timestamp"],
        ["listings","id","Integer","PRIMARY KEY, AUTO","Crop listing identifier"],
        ["listings","farmer_id","Integer","FK → users.id","Farmer who created listing"],
        ["listings","crop_name","String(100)","NOT NULL","Name of crop being sold"],
        ["listings","quantity_kg","Float","NOT NULL","Quantity available in kg"],
        ["listings","price_per_kg","Float","NOT NULL","Asking price per kg (INR)"],
        ["listings","status","String(20)","DEFAULT 'available'","available/purchased/confirmed/completed"],
        ["purchases","id","Integer","PRIMARY KEY, AUTO","Purchase transaction identifier"],
        ["purchases","listing_id","Integer","FK → listings.id","Associated listing"],
        ["purchases","buyer_id","Integer","FK → users.id","Buyer who initiated purchase"],
        ["purchases","quantity_kg","Float","NOT NULL","Quantity purchased"],
        ["purchases","total_price","Float","NOT NULL","Total transaction value (INR)"],
        ["messages","id","Integer","PRIMARY KEY, AUTO","Message identifier"],
        ["messages","listing_id","Integer","FK → listings.id","Associated listing/negotiation"],
        ["messages","sender_id","Integer","FK → users.id","Message sender"],
        ["messages","content","Text","NOT NULL","Message body"],
        ["messages","timestamp","DateTime","DEFAULT NOW()","Message timestamp"],
        ["finance_records","id","Integer","PRIMARY KEY, AUTO","Financial record identifier"],
        ["finance_records","user_id","Integer","FK → users.id","Associated user (farmer)"],
        ["finance_records","type","String(10)","NOT NULL","'income' or 'expense'"],
        ["finance_records","amount","Float","NOT NULL","Amount in INR"],
        ["finance_records","description","String(255)","","Description of transaction"],
        ["calendar_tasks","id","Integer","PRIMARY KEY, AUTO","Calendar task identifier"],
        ["calendar_tasks","user_id","Integer","FK → users.id","Associated farmer"],
        ["calendar_tasks","crop_name","String(100)","NOT NULL","Crop associated with task"],
        ["calendar_tasks","task_description","Text","NOT NULL","Auto-generated task description"],
        ["calendar_tasks","due_date","Date","NOT NULL","Task due date"],
        ["calendar_tasks","completed","Boolean","DEFAULT False","Task completion status"],
    ],
    "Table 5.5 KrishiMitra Database Schema"
)

br(doc)
sec(doc,"References for Chapter 5")
for r in [
    "[1] Pressman, R. S. and Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach, 9th ed. McGraw-Hill Education.",
    "[2] Fowler, M. (2002). Patterns of Enterprise Application Architecture. Addison-Wesley Professional.",
    "[3] Google LLC (2024). Flutter Documentation. Available: https://flutter.dev/docs",
    "[4] Freeman, E., Robson, E., Bates, B. and Sierra, K. (2004). Head First Design Patterns. O'Reilly Media.",
    "[5] IETF (2022). RFC 7230: Hypertext Transfer Protocol (HTTP/1.1). Internet Engineering Task Force.",
    "[6] Microsoft (2023). Three-tier application architecture. Azure Architecture Center. Available: https://learn.microsoft.com/en-us/azure/architecture",
    "[7] Jacobson, I., Booch, G. and Rumbaugh, J. (1999). The Unified Software Development Process. Addison-Wesley.",
    "[8] Yourdon, E. (1989). Modern Structured Analysis. Prentice Hall.",
    "[9] Kamilaris, A. and Prenafeta-Boldu, F. X. (2018). Deep learning in agriculture: A survey. Computers and Electronics in Agriculture, 147, pp. 70–90.",
    "[10] Fielding, R. T. (2000). Architectural Styles and the Design of Network-based Software Architectures. Doctoral dissertation, University of California, Irvine.",
]:
    body(doc,r)

doc.save("KrishiMitra_Ch5.docx")
print("Chapter 5 saved to KrishiMitra_Ch5.docx")
