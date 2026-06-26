import urllib.request
import base64
import time
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def download_mermaid_ink(mermaid_str, filename):
    encoded = base64.b64encode(mermaid_str.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}'
    
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    print(f'Downloading {filename}...')
    try:
        with urllib.request.urlopen(req, timeout=30) as response, open(filename, 'wb') as out_file:
            out_file.write(response.read())
        print(f'Success: {filename}')
        return True
    except Exception as e:
        print(f'Failed {filename}: {e}')
        return False

# Diagrams
architecture_code = """
graph TD
    subgraph Mobile Client
        UI[Flutter App]
    end
    subgraph Backend Server
        API[Flask REST API]
        DB[(SQLite Database)]
        API --> DB
    end
    subgraph Intelligence Layer
        RN[ResNet50 Model]
        KNN[KNN Soil Engine]
        LSTM[LSTM Forecaster]
        LLM[Groq LLaMA 3.3]
    end
    UI <-->|JSON / HTTPS| API
    API <-->|Inference| RN
    API <-->|Analysis| KNN
    API <-->|Prediction| LSTM
    API <-->|API Calls| LLM
"""

er_code = """
erDiagram
    USER ||--o{ LISTING : "creates"
    USER ||--o{ PURCHASE : "makes"
    USER ||--o{ MESSAGE : "sends"
    USER ||--o{ TRANSACTION : "has"
    USER ||--o{ CALENDARTASK : "manages"
    USER ||--o{ ECOMMERCEORDER : "places"
    LISTING ||--o{ PURCHASE : "contains"
    PURCHASE ||--o{ MESSAGE : "tracks"
"""

dfd_code = """
graph LR
    F((Farmer)) -->|Images, Soil Data| System[KrishiMitra System]
    B((Buyer)) -->|Search, Bids| System
    System -->|Recommendations| F
    System -->|Listings| B
    System <-->|Prices| Agmarknet[Agmarknet API]
    System <-->|Advisory| Groq[Groq API]
"""

flowchart_code = """
graph TD
    A([Start]) --> B[Capture Leaf Image]
    B --> C[Compress & Encode Base64]
    C --> D[Transmit via HTTPS POST]
    D --> E{Backend Flask API}
    E --> F[Decode Image Tensor]
    F --> G[ResNet50 Inference]
    G --> H[Extract Feature Maps]
    H --> I[Apply Grad-CAM]
    I --> J[Overlay JET Colormap]
    J --> K[Return JSON + Image URL]
    K --> L[Render Heatmap]
    L --> M([End])
"""

diagrams = [
    (architecture_code, "real_arch.png", "System Architecture Design", "1.4 Proposed Approach", 1),
    (dfd_code, "real_dfd.png", "Data Flow Diagram (Level 0)", "5.3.2 Data Flow Diagrams (DFD)", 2),
    (er_code, "real_er.png", "Entity-Relationship (ER) Diagram", "5.4 Database Design and Data Dictionary", 3),
    (flowchart_code, "real_flow.png", "Flowchart: Disease Detection Pipeline", "6.2.1 Deep Learning Vision Module", 4)
]

for code, filename, caption, target_text, fig_num in diagrams:
    if not os.path.exists(filename):
        success = download_mermaid_ink(code, filename)
        if success:
            time.sleep(5) # rate limit prevention

# Open the base document (without the basic appendices)
doc_path = "KrishiMitra_Final_Report_Massive.docx"
print(f"Opening {doc_path}...")
try:
    doc = Document(doc_path)
except Exception as e:
    print(f"Error opening doc: {e}")
    exit(1)

for code, filename, caption, target_text, fig_num in diagrams:
    if not os.path.exists(filename):
        continue
        
    for i, p in enumerate(doc.paragraphs):
        if target_text in p.text:
            print(f"Found target: {target_text}")
            # Insert after this paragraph (by inserting before paragraph i+2 to leave some space)
            # Make sure we don't go out of bounds
            insert_idx = min(i + 2, len(doc.paragraphs) - 1)
            target_p = doc.paragraphs[insert_idx]
            
            # Add image paragraph
            new_p = target_p.insert_paragraph_before()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_p.add_run()
            run.add_picture(filename, width=Inches(5.0))
            
            # Add caption paragraph
            cap_p = target_p.insert_paragraph_before()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_r = cap_p.add_run(f"Figure {fig_num}: {caption}")
            cap_r.font.size = Pt(12)
            cap_r.font.name = "Times New Roman"
            cap_r.bold = True
            
            # Add a blank line after
            target_p.insert_paragraph_before("")
            break

out_name = "KrishiMitra_Final_Report_Massive_Correct_Diagrams.docx"
doc.save(out_name)
print(f"Successfully inserted correct diagrams into {out_name}!")
