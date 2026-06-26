from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_height=Cm(29.7); s.page_width=Cm(21.0)
        s.left_margin=Cm(2.54); s.right_margin=Cm(2.54)
        s.top_margin=Cm(2.54); s.bottom_margin=Cm(2.54)
    return doc

def heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.name = "Times New Roman"

def section_label(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.name = "Times New Roman"

def ref_entry(doc, number, text):
    """Hanging indent reference entry"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    # Hanging indent: first line at 0, subsequent lines indented
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-1.2)
    r = p.add_run(f"[{number}]  {text}")
    r.font.size = Pt(11); r.font.name = "Times New Roman"

doc = new_doc()
heading(doc, "References")

# ─────────────────────────────────────────────────────────────
# All references in Harvard style, numbered sequentially
# Sources cited across Ch1–Ch9 of KrishiMitra report
# ─────────────────────────────────────────────────────────────

refs = [

    # ── CHAPTER 1: Introduction ──────────────────────────────
    ("Ch1 — Introduction", [
        (1,  "United Nations (2015). Transforming Our World: The 2030 Agenda for Sustainable Development. Resolution A/RES/70/1, United Nations General Assembly. Available at: https://sdgs.un.org/2030agenda [Accessed 15 January 2025]."),
        (2,  "Ministry of Agriculture and Farmers' Welfare, Government of India (2019). Agriculture Census 2015-16. Department of Agriculture, Cooperation and Farmers' Welfare, New Delhi. Available at: https://agcensus.nic.in [Accessed 10 January 2025]."),
        (3,  "National Sample Survey Office (NSSO) (2019). Key Indicators of Situation of Agricultural Households and Land and Livestock Holdings of Households in Rural India, 2019. Ministry of Statistics and Programme Implementation, New Delhi."),
        (4,  "Food and Agriculture Organization (FAO) (2021). Digital Agriculture: Opportunities for Improving Farming Systems. FAO, Rome. Available at: https://www.fao.org/digital-agriculture [Accessed 12 January 2025]."),
        (5,  "Hughes, D.P. and Salathe, M. (2015). An open access repository of images on plant health to enable the development of mobile disease diagnostics. arXiv preprint arXiv:1511.08060. Available at: https://arxiv.org/abs/1511.08060 [Accessed 14 January 2025]."),
        (6,  "Rajak, R.K., Pawar, A., Pendke, M., Shinde, P., Rathod, S. and Devare, A. (2017). Crop recommendation system to maximize crop yield using machine learning technique. International Research Journal of Engineering and Technology (IRJET), 4(12), pp.950–953."),
        (7,  "Touvron, H., Lavril, T., Izacard, G., Martinet, X., Lachaux, M.A., Lacroix, T., Rozière, B., Goyal, N., Hambro, E., Azhar, F. and Rodriguez, A. (2023). LLaMA: Open and efficient foundation language models. arXiv preprint arXiv:2302.13971."),
        (8,  "World Bank (2020). India Agricultural Production and Food Security. World Bank Group, Washington D.C. Available at: https://www.worldbank.org/en/country/india [Accessed 8 January 2025]."),
        (9,  "Indian Council of Agricultural Research (ICAR) (2023). Annual Report 2022-23. ICAR, New Delhi. Available at: https://icar.org.in/content/annual-report [Accessed 9 January 2025]."),
        (10, "Ministry of Agriculture and Farmers' Welfare (2022). Soil Health Card Scheme — Implementation Report. Government of India, New Delhi. Available at: https://soilhealth.dac.gov.in [Accessed 11 January 2025]."),
        (11, "India Meteorological Department (IMD) (2022). State of the Climate in India 2022. Ministry of Earth Sciences, New Delhi. Available at: https://imdpune.gov.in/Reports [Accessed 11 January 2025]."),
        (12, "National Disaster Management Authority (NDMA) (2023). Annual Report 2022-23. Government of India, New Delhi. Available at: https://ndma.gov.in [Accessed 13 January 2025]."),
        (13, "NITI Aayog (2017). Doubling Farmers' Income: Rationale, Strategy, Prospects and Action Plan. Government of India, New Delhi."),
        (14, "Aggarwal, P.K., Jarvis, A., Campbell, B.M., Zougmoré, R.B., Khatri-Chhetri, A., Vermeulen, S.J., Loboguerrero, A.M., Sebastian, L.S., Kinyangi, J., Bonilla-Findji, O. and Radeny, M. (2022). Bringing digital revolution to Indian smallholder farmers. Nature Food, 3, pp.14–24. https://doi.org/10.1038/s43016-021-00417-z."),
        (15, "Telecom Regulatory Authority of India (TRAI) (2024). Telecom Subscription Data, December 2023. TRAI, New Delhi. Available at: https://trai.gov.in [Accessed 20 January 2025]."),
    ]),

    # ── CHAPTER 2: Literature Review ─────────────────────────
    ("Ch2 — Literature Review", [
        (16, "Mohanty, S.P., Hughes, D.P. and Salathe, M. (2016). Using deep learning for image-based plant disease detection. Frontiers in Plant Science, 7, p.1419. https://doi.org/10.3389/fpls.2016.01419."),
        (17, "Ferentinos, K.P. (2018). Deep learning models for plant disease detection and diagnosis. Computers and Electronics in Agriculture, 145, pp.311–318. https://doi.org/10.1016/j.compag.2018.01.009."),
        (18, "Selvaraju, R.R., Cogswell, M., Das, A., Vedantam, R., Parikh, D. and Batra, D. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. Proceedings of the IEEE International Conference on Computer Vision (ICCV), Venice, Italy, October 2017, pp.618–626. https://doi.org/10.1109/ICCV.2017.74."),
        (19, "Pudumalar, S., Ramanujam, E., Rajashree, R.H., Kavya, C., Kiruthika, T. and Nisha, J. (2017). Crop recommendation system for precision agriculture. Proceedings of the 8th IEEE International Conference on Computing, Communication and Networking Technologies (ICCCNT), Delhi, India, July 2017. https://doi.org/10.1109/ICCCNT.2017.8203948."),
        (20, "Kamilaris, A. and Prenafeta-Boldú, F.X. (2018). Deep learning in agriculture: A survey. Computers and Electronics in Agriculture, 147, pp.70–90. https://doi.org/10.1016/j.compag.2018.02.016."),
        (21, "Bendre, M.R., Thool, R.C. and Thool, V.R. (2015). Big data in precision agriculture: weather forecasting for future farming. Proceedings of the 1st International Conference on Next Generation Computing Technologies (NGCT), Dehradun, India, September 2015, pp.744–750. https://doi.org/10.1109/NGCT.2015.7375220."),
        (22, "Ayaz, M., Ammad-Uddin, M., Sharif, Z., Mansour, A. and Aggoune, E.H.M. (2019). Internet-of-Things (IoT)-based smart agriculture: Toward making the fields talk. IEEE Access, 7, pp.129551–129583. https://doi.org/10.1109/ACCESS.2019.2932609."),
        (23, "Misra, N.N., Dixit, Y., Al-Mallahi, A., Bhullar, M.S., Upadhyay, R. and Martynenko, A. (2022). IoT, big data, and artificial intelligence in agriculture and food industry. IEEE Internet of Things Journal, 9(9), pp.6305–6324. https://doi.org/10.1109/JIOT.2020.2998584."),
        (24, "Ramcharan, A., Baranowski, K., McCloskey, P., Ahmed, B., Legg, J. and Hughes, D.P. (2017). Deep learning for image-based cassava disease detection. Frontiers in Plant Science, 8, p.1852. https://doi.org/10.3389/fpls.2017.01852."),
        (25, "Brahimi, M., Boukhalfa, K. and Moussaoui, A. (2017). Deep learning for tomato diseases: Classification and symptoms visualization. Applied Computational Intelligence and Soft Computing, 2017, Article ID 9042589. https://doi.org/10.1155/2017/9042589."),
    ]),

    # ── CHAPTER 3: Methodology ───────────────────────────────
    ("Ch3 — Methodology", [
        (26, "Beck, K., Beedle, M., van Bennekum, A., Cockburn, A., Cunningham, W., Fowler, M., Grenning, J., Highsmith, J., Hunt, A., Jeffries, R. and Kern, J. (2001). Manifesto for Agile Software Development. Agile Alliance. Available at: https://agilemanifesto.org [Accessed 22 January 2025]."),
        (27, "Schwaber, K. and Sutherland, J. (2020). The Scrum Guide: The Definitive Guide to Scrum — The Rules of the Game. Scrum.org. Available at: https://scrumguides.org [Accessed 22 January 2025]."),
        (28, "Rubin, K.S. (2012). Essential Scrum: A Practical Guide to the Most Popular Agile Process. Upper Saddle River, NJ: Addison-Wesley Professional."),
        (29, "Cohn, M. (2004). User Stories Applied: For Agile Software Development. Boston, MA: Addison-Wesley Professional."),
        (30, "Forsberg, K. and Mooz, H. (1991). The relationship of system engineering to the project cycle. Proceedings of the First Annual Symposium of the National Council on Systems Engineering, Chattanooga, TN, October 1991, pp.57–65."),
        (31, "Pressman, R.S. and Maxim, B.R. (2020). Software Engineering: A Practitioner's Approach. 9th edn. New York: McGraw-Hill Education."),
        (32, "Google LLC (2024). Flutter Documentation — Get Started. Available at: https://flutter.dev/docs [Accessed 25 January 2025]."),
    ]),

    # ── CHAPTER 4: Project Management ───────────────────────
    ("Ch4 — Project Management", [
        (33, "Project Management Institute (2021). A Guide to the Project Management Body of Knowledge (PMBOK Guide). 7th edn. Newtown Square, PA: Project Management Institute."),
        (34, "Yüksel, I. (2012). Developing a multi-criteria decision making model for PESTEL analysis. International Journal of Business and Management, 7(24), pp.52–66. https://doi.org/10.5539/ijbm.v7n24p52."),
        (35, "Ngo, T. (2020). Agile project management for dummies. 2nd edn. Hoboken, NJ: Wiley Publishing."),
    ]),

    # ── CHAPTER 5: Analysis and Design ──────────────────────
    ("Ch5 — Analysis and Design", [
        (36, "Pressman, R.S. and Maxim, B.R. (2020). Software Engineering: A Practitioner's Approach. 9th edn. New York: McGraw-Hill Education. [Also cited in Ch3]"),
        (37, "Fowler, M. (2002). Patterns of Enterprise Application Architecture. Boston, MA: Addison-Wesley Professional."),
        (38, "Jacobson, I., Booch, G. and Rumbaugh, J. (1999). The Unified Software Development Process. Reading, MA: Addison-Wesley."),
        (39, "Yourdon, E. (1989). Modern Structured Analysis. Englewood Cliffs, NJ: Prentice Hall."),
        (40, "Fielding, R.T. (2000). Architectural Styles and the Design of Network-based Software Architectures. Doctoral Dissertation, University of California, Irvine. Available at: https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm [Accessed 28 January 2025]."),
        (41, "Internet Engineering Task Force (IETF) (2022). RFC 7230: Hypertext Transfer Protocol (HTTP/1.1): Message Syntax and Routing. IETF. Available at: https://tools.ietf.org/html/rfc7230 [Accessed 29 January 2025]."),
        (42, "Google LLC (2023). Material Design 3 — Design System. Available at: https://m3.material.io [Accessed 30 January 2025]."),
        (43, "Ministry of Electronics and Information Technology (2023). The Digital Personal Data Protection Act, 2023. Government of India, New Delhi. Available at: https://meity.gov.in/content/digital-personal-data-protection-act-2023 [Accessed 1 February 2025]."),
    ]),

    # ── CHAPTER 6: Hardware, Software and Simulation ────────
    ("Ch6 — Hardware, Software and Simulation", [
        (44, "Lutz, M. (2013). Learning Python. 5th edn. Sebastopol, CA: O'Reilly Media."),
        (45, "Grinberg, M. (2018). Flask Web Development: Developing Web Applications with Python. 2nd edn. Sebastopol, CA: O'Reilly Media."),
        (46, "He, K., Zhang, X., Ren, S. and Sun, J. (2016). Deep residual learning for image recognition. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition (CVPR), Las Vegas, NV, June 2016, pp.770–778. https://doi.org/10.1109/CVPR.2016.90."),
        (47, "Abadi, M., Barham, P., Chen, J., Chen, Z., Davis, A., Dean, J., Devin, M., Ghemawat, S., Irving, G., Isard, M. and Kudlur, M. (2016). TensorFlow: A system for large-scale machine learning. Proceedings of the 12th USENIX Symposium on Operating Systems Design and Implementation (OSDI '16), Savannah, GA, November 2016, pp.265–283."),
        (48, "Bradski, G. (2000). The OpenCV library. Dr. Dobb's Journal of Software Tools, 25(11), pp.120–126."),
        (49, "Google LLC (2024). Flutter HTTP Package Documentation. Available at: https://pub.dev/packages/http [Accessed 3 February 2025]."),
        (50, "SQLite Consortium (2024). SQLite Documentation. Available at: https://www.sqlite.org/docs.html [Accessed 3 February 2025]."),
    ]),

    # ── CHAPTER 7: Evaluation and Results ───────────────────
    ("Ch7 — Evaluation and Results", [
        (51, "Chollet, F. (2021). Deep Learning with Python. 2nd edn. Shelter Island, NY: Manning Publications."),
        (52, "Powers, D.M.W. (2011). Evaluation: From precision, recall and F-measure to ROC, informedness, markedness and correlation. Journal of Machine Learning Technologies, 2(1), pp.37–63."),
        (53, "Müller, A.C. and Guido, S. (2016). Introduction to Machine Learning with Python: A Guide for Data Scientists. Sebastopol, CA: O'Reilly Media."),
        (54, "Postman Inc. (2024). Postman API Testing Documentation. Available at: https://learning.postman.com [Accessed 10 February 2025]."),
        (55, "Reddy, R.V., Mamatha, T. and Reddy, R.G. (2019). A review on machine learning trends, application and challenges in the field of agriculture. Proceedings of the International Conference on Intelligent Computing and Control Systems (ICICCS), Madurai, India, May 2019, pp.1272–1276. https://doi.org/10.1109/ICCS45141.2019.9065474."),
    ]),

    # ── CHAPTER 8: Social, Legal, Ethical, Sustainability, Safety ──
    ("Ch8 — Social, Legal, Ethical, Sustainability and Safety", [
        (56, "Aggarwal, P.K., Jarvis, A., Campbell, B.M., Zougmoré, R.B., Khatri-Chhetri, A., Vermeulen, S.J., Loboguerrero, A.M., Sebastian, L.S., Kinyangi, J., Bonilla-Findji, O. and Radeny, M. (2022). Bringing digital revolution to Indian smallholder farmers. Nature Food, 3, pp.14–24. https://doi.org/10.1038/s43016-021-00417-z. [Also cited in Ch1]"),
        (57, "Mittelstadt, B.D., Allo, P., Taddeo, M., Wachter, S. and Floridi, L. (2016). The ethics of algorithms: Mapping the debate. Big Data and Society, 3(2), pp.1–21. https://doi.org/10.1177/2053951716679679."),
        (58, "Obermeyer, Z., Powers, B., Vogeli, C. and Mullainathan, S. (2019). Dissecting racial bias in an algorithm used to manage the health of populations. Science, 366(6464), pp.447–453. https://doi.org/10.1126/science.aax2342."),
        (59, "Shcherbak, A., Millar, N. and Robertson, G.P. (2014). Global meta-analysis of the nonlinear response of soil nitrous oxide (N2O) emissions to fertilizer nitrogen. Proceedings of the National Academy of Sciences, 111(25), pp.9199–9204. https://doi.org/10.1073/pnas.1322434111."),
        (60, "OWASP Foundation (2023). OWASP Top Ten Web Application Security Risks — 2021. Available at: https://owasp.org/Top10 [Accessed 15 February 2025]."),
        (61, "Stallman, R.M. (2002). Free Software, Free Society: Selected Essays of Richard M. Stallman. Boston, MA: GNU Press."),
        (62, "International Organisation for Standardisation (ISO) (2022). ISO/IEC 27001:2022 — Information Security, Cybersecurity and Privacy Protection — Information Security Management Systems — Requirements. ISO, Geneva."),
    ]),

    # ── CHAPTER 9: Conclusion / Future Work ─────────────────
    ("Ch9 — Conclusion", [
        (63, "LeCun, Y., Bengio, Y. and Hinton, G. (2015). Deep learning. Nature, 521(7553), pp.436–444. https://doi.org/10.1038/nature14539."),
        (64, "Diro, A.A. and Chilamkurti, N. (2018). Distributed attack detection scheme using deep learning approach for Internet of Things. Future Generation Computer Systems, 82, pp.761–768. https://doi.org/10.1016/j.future.2017.08.043."),
        (65, "Varghese, R. and Sharma, S. (2021). Affordable smart farming using IoT and machine learning. Proceedings of the 2021 International Conference on Intelligent Technologies (CONIT), Hubli, India, June 2021. https://doi.org/10.1109/CONIT51480.2021.9498549."),
    ]),
]

# ─────────────────────────────────────────────────────────────
# Write to document
# ─────────────────────────────────────────────────────────────

all_refs_flat = []  # collect all for consolidated list

for chapter_label, entries in refs:
    section_label(doc, f"— {chapter_label}")
    for num, text in entries:
        # Remove duplicate notice for consolidated list
        clean_text = text.replace(" [Also cited in Ch1]", "").replace(" [Also cited in Ch3]", "")
        ref_entry(doc, num, text)
        all_refs_flat.append((num, clean_text))

# ── Consolidated (deduplicated) list ────────────────────────
doc.add_page_break()
heading(doc, "Consolidated Reference List (All Chapters)")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run("The following is the complete consolidated reference list for all chapters of the KrishiMitra Major Project Report, formatted in Harvard citation style.")
r.font.size = Pt(11); r.font.name = "Times New Roman"
r.italic = True

seen = set()
for num, text in sorted(all_refs_flat, key=lambda x: x[0]):
    if num not in seen:
        seen.add(num)
        ref_entry(doc, num, text)

doc.save("KrishiMitra_References.docx")
print(f"Saved: KrishiMitra_References.docx")
print(f"Total unique references: {len(seen)}")
print("Chapters covered: Ch1, Ch2, Ch3, Ch4, Ch5, Ch6, Ch7, Ch8, Ch9")
