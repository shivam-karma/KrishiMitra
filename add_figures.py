from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def get_doc(filename="KrishiMitra_Final_Report_Massive.docx"):
    return Document(filename)

def para(doc, text, size=12, bold=False, align="JUSTIFY"):
    p = doc.add_paragraph()
    if   align == "CENTER":  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "LEFT":    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "RIGHT":   p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    from docx.shared import Pt as Pt2
    pf.line_spacing = Pt2(21)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    return p

def chap(doc, text):
    para(doc, text, size=16, bold=True, align="CENTER")

def sec(doc, text):
    para(doc, text, size=14, bold=True, align="LEFT")

def blank(doc):
    para(doc, "", size=12, align="LEFT")

def add_image_with_caption(doc, img_path, caption, fig_num):
    if not os.path.exists(img_path):
        return
    
    # Add Image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Add picture, scale width to fit reasonably
    run.add_picture(img_path, width=Inches(3.5))
    
    # Add Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cap = p_cap.add_run(f"Figure C.{fig_num}: {caption}")
    r_cap.font.size = Pt(12)
    r_cap.font.name = "Times New Roman"
    r_cap.bold = True
    blank(doc)

doc = get_doc()

doc.add_page_break()

sec(doc, "Appendix C: Application User Interfaces (Simulation Screenshots)")
blank(doc)
para(doc, "The following figures demonstrate the implemented UI screens for the various intelligent modules of the KrishiMitra mobile application, developed using the Flutter framework.", align="LEFT")
blank(doc)

screenshots = [
    ("login_role_selection", "User Authentication and Role Selection Interface"),
    ("krishi_dashboard", "Primary Farmer Dashboard (Navigation Hub)"),
    ("krishi_premium_dashboard", "Premium Analytics Dashboard"),
    ("disease_detection", "ResNet50 Leaf Disease Detection with Grad-CAM"),
    ("soil_analysis_recommendation", "KNN Soil Recommendation and Health Score"),
    ("market_trends", "LSTM Mandi Price Trend Forecasting"),
    ("ai_chatbot", "Groq LLaMA 3.3 Conversational Advisory Interface"),
    ("buyer_marketplace_feed", "P2P Crop Marketplace - Buyer Feed"),
    ("crop_listing_details", "Marketplace Crop Listing Details"),
    ("p2p_marketplace_chat", "Real-Time Negotiation Chat Interface"),
    ("ai_farming_calendar", "Automated Farming Task Calendar"),
    ("government_schemes_hub", "Government Subsidies and Schemes Hub"),
    ("user_profile_analytics", "User Profile and Financial Analytics")
]

fig_idx = 1
for folder, desc in screenshots:
    path = os.path.join(r"C:\Users\shiva\Downloads\stitch\stitch", folder, "screen.png")
    if os.path.exists(path):
        add_image_with_caption(doc, path, desc, fig_idx)
        fig_idx += 1

doc.save("KrishiMitra_Final_Report_Massive_With_Figures.docx")
print("Added Appendix C with Application Screenshots. Saved as KrishiMitra_Final_Report_Massive_With_Figures.docx")
