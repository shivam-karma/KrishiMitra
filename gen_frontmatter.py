from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_height=Cm(29.7); s.page_width=Cm(21.0)
        s.left_margin=Cm(2.54); s.right_margin=Cm(2.54)
        s.top_margin=Cm(2.54); s.bottom_margin=Cm(2.54)
    return doc

def heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.name = "Times New Roman"

def blank(doc):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = Pt(12)

def make_list_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=3)
    t.style = "Table Grid"
    # header row
    hr = t.rows[0]
    hdata = headers
    for i, h in enumerate(hdata):
        c = hr.cells[i]; c.text = h
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True; run.font.size = Pt(11); run.font.name = "Times New Roman"
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864'); shd.set(qn('w:val'), 'clear'); tcPr.append(shd)
        for run in p.runs: run.font.color.rgb = RGBColor(255,255,255)
    # data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        fill = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]
        for ci, val in enumerate(row):
            c = tr.cells[ci]; c.text = str(val)
            p2 = c.paragraphs[0]; p2.alignment = aligns[ci]
            p2.paragraph_format.space_before = Pt(3)
            p2.paragraph_format.space_after = Pt(3)
            for run in p2.runs:
                run.font.size = Pt(11); run.font.name = "Times New Roman"
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill); shd.set(qn('w:val'), 'clear'); tcPr.append(shd)
    # set column widths
    widths = [Cm(2.5), Cm(12.0), Cm(2.5)]
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    blank(doc)

doc = new_doc()

# ═══════════════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════════════
heading(doc, "List of Figures")

figures = [
    ("1.1",  "Comparative feature matrix of KrishiMitra vs existing agricultural apps", "8"),
    ("1.2",  "KrishiMitra alignment with UN Sustainable Development Goals (SDGs)", "11"),
    ("3.1",  "Agile Scrum Methodology applied to KrishiMitra development", "20"),
    ("3.2",  "V-Model Verification and Validation phases applied to KrishiMitra", "22"),
    ("5.1",  "Functional Block Diagram of KrishiMitra System", "34"),
    ("5.2",  "System Flowchart for KrishiMitra Application", "36"),
    ("5.3",  "Three-Tier System Architecture of KrishiMitra", "39"),
    ("5.4",  "Use Case Diagram for KrishiMitra", "41"),
    ("5.5",  "Level 0 Data Flow Diagram (Context Diagram) for KrishiMitra", "43"),
    ("5.6",  "Level 1 Data Flow Diagram for KrishiMitra", "44"),
    ("5.7",  "Request-Response Communication Model for KrishiMitra", "46"),
    ("6.1",  "CNN Disease Detection Pipeline — ResNet50 with Grad-CAM overlay", "53"),
    ("6.2",  "KNN Soil Analysis Algorithm — Feature space and k=7 nearest neighbour selection", "57"),
    ("6.3",  "Flutter Disease Detection Screen — Grad-CAM heatmap rendered on leaf image", "60"),
    ("7.1",  "CNN Model Training and Validation Accuracy/Loss Curves (25 epochs)", "68"),
    ("7.2",  "Confusion Matrix — CNN Disease Detection (3 classes, 2180 test images)", "69"),
    ("7.3",  "API Response Time Distribution across all 20+ endpoints (Postman benchmark)", "72"),
    ("7.4",  "KNN Crop Recommendation Accuracy by Crop Type (Top-1 and Top-3)", "73"),
]

make_list_table(doc, ["Figure No.", "Caption", "Page No."], figures)

doc.add_page_break()

# ═══════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════
heading(doc, "List of Tables")

tables = [
    ("2.1",  "Summary of Literature Reviews", "15"),
    ("3.1",  "Mapping KrishiMitra Development Stages to Agile Scrum Phases", "21"),
    ("4.1",  "Project Planning Timeline", "26"),
    ("4.2",  "Project Implementation Timeline (Sprint Schedule)", "27"),
    ("4.3",  "PESTLE Risk Analysis for KrishiMitra", "29"),
    ("4.4",  "Estimated Project Budget for KrishiMitra", "31"),
    ("5.1",  "KrishiMitra System Requirements Specification", "34"),
    ("5.2",  "Comparison of Mobile Frontend Frameworks", "37"),
    ("5.3",  "Comparison of Backend Web Frameworks", "38"),
    ("5.4",  "Standards and Protocols Applied in KrishiMitra", "40"),
    ("5.5",  "KrishiMitra Database Schema (7 Tables)", "47"),
    ("6.1",  "Development and Testing Hardware Configuration", "51"),
    ("6.2",  "Software Development Tools Used in KrishiMitra", "52"),
    ("6.3",  "Simulation and Testing Tools", "63"),
    ("7.1",  "KrishiMitra Test Plan and Results (22 Test Cases)", "67"),
    ("7.2",  "CNN Disease Detection Model Performance — Precision, Recall, F1-Score", "69"),
    ("7.3",  "KNN Soil Crop Recommendation Performance Metrics", "71"),
    ("7.4",  "API Response Time Benchmarks (All Endpoints)", "72"),
]

make_list_table(doc, ["Table No.", "Caption", "Page No."], tables)

doc.add_page_break()

# ═══════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════
heading(doc, "List of Abbreviations")

abbreviations = [
    ("ADC",    "Analog to Digital Converter"),
    ("AI",     "Artificial Intelligence"),
    ("API",    "Application Programming Interface"),
    ("APK",    "Android Package Kit"),
    ("ARM",    "Advanced RISC Machine"),
    ("BERT",   "Bidirectional Encoder Representations from Transformers"),
    ("CNN",    "Convolutional Neural Network"),
    ("CORS",   "Cross-Origin Resource Sharing"),
    ("CPU",    "Central Processing Unit"),
    ("CRUD",   "Create, Read, Update, Delete"),
    ("CSV",    "Comma-Separated Values"),
    ("DPDPA",  "Digital Personal Data Protection Act"),
    ("DFD",    "Data Flow Diagram"),
    ("FAO",    "Food and Agriculture Organisation"),
    ("GDP",    "Gross Domestic Product"),
    ("GPU",    "Graphics Processing Unit"),
    ("Grad-CAM","Gradient-weighted Class Activation Mapping"),
    ("HTTP",   "Hypertext Transfer Protocol"),
    ("HTTPS",  "Hypertext Transfer Protocol Secure"),
    ("ICAR",   "Indian Council of Agricultural Research"),
    ("ICT",    "Information and Communication Technology"),
    ("IDE",    "Integrated Development Environment"),
    ("IMD",    "India Meteorological Department"),
    ("IoT",    "Internet of Things"),
    ("IRJET",  "International Research Journal of Engineering and Technology"),
    ("JSON",   "JavaScript Object Notation"),
    ("JWT",    "JSON Web Token"),
    ("KNN",    "K-Nearest Neighbours"),
    ("LLM",    "Large Language Model"),
    ("LSTM",   "Long Short-Term Memory"),
    ("MCA",    "Master of Computer Applications"),
    ("ML",     "Machine Learning"),
    ("MQTT",   "Message Queuing Telemetry Transport"),
    ("MSE",    "Mean Squared Error"),
    ("NLP",    "Natural Language Processing"),
    ("NPK",    "Nitrogen, Phosphorus, Potassium"),
    ("NSSO",   "National Sample Survey Office"),
    ("ORM",    "Object-Relational Mapping"),
    ("OTP",    "One-Time Password"),
    ("P2P",    "Peer-to-Peer"),
    ("PESTLE", "Political, Economic, Social, Technological, Legal, Environmental"),
    ("pH",     "Potential of Hydrogen (soil acidity measure)"),
    ("REST",   "Representational State Transfer"),
    ("RMSE",   "Root Mean Squared Error"),
    ("SDG",    "Sustainable Development Goal"),
    ("SMS",    "Short Message Service"),
    ("SQL",    "Structured Query Language"),
    ("SQLite", "Structured Query Language Lite (lightweight database engine)"),
    ("TLS",    "Transport Layer Security"),
    ("UI",     "User Interface"),
    ("UML",    "Unified Modelling Language"),
    ("UPI",    "Unified Payments Interface"),
    ("URL",    "Uniform Resource Locator"),
    ("UX",     "User Experience"),
    ("VCS",    "Version Control System"),
    ("XAI",    "Explainable Artificial Intelligence"),
    ("XML",    "Extensible Markup Language"),
]

# Two-column abbreviation layout using a table
t = doc.add_table(rows=len(abbreviations), cols=2)
t.style = "Table Grid"
for ri, (abbr, full) in enumerate(abbreviations):
    fill = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
    # Abbreviation cell
    c1 = t.rows[ri].cells[0]; c1.text = abbr
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_before = Pt(2); p1.paragraph_format.space_after = Pt(2)
    for run in p1.runs:
        run.bold = True; run.font.size = Pt(11); run.font.name = "Times New Roman"
    tc1 = c1._tc; tcPr1 = tc1.get_or_add_tcPr()
    shd1 = OxmlElement('w:shd'); shd1.set(qn('w:fill'), fill); shd1.set(qn('w:val'), 'clear'); tcPr1.append(shd1)
    # Full form cell
    c2 = t.rows[ri].cells[1]; c2.text = full
    p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
    for run in p2.runs:
        run.font.size = Pt(11); run.font.name = "Times New Roman"
    tc2 = c2._tc; tcPr2 = tc2.get_or_add_tcPr()
    shd2 = OxmlElement('w:shd'); shd2.set(qn('w:fill'), fill); shd2.set(qn('w:val'), 'clear'); tcPr2.append(shd2)

for row in t.rows:
    row.cells[0].width = Cm(4.0)
    row.cells[1].width = Cm(13.0)

doc.save("KrishiMitra_FrontMatter.docx")
print("Saved: KrishiMitra_FrontMatter.docx")
print(f"  - List of Figures: {len(figures)} entries")
print(f"  - List of Tables:  {len(tables)} entries")
print(f"  - Abbreviations:   {len(abbreviations)} entries")
