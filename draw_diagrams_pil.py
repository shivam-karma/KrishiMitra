from PIL import Image, ImageDraw, ImageFont
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_diagram(filename, title, boxes):
    # Create a white image
    img = Image.new('RGB', (800, 600), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    
    # Title
    d.text((50, 30), title, fill=(0, 0, 0))
    
    # Draw boxes
    # boxes format: [(x, y, w, h, text)]
    for (x, y, w, h, text) in boxes:
        # Draw rectangle
        d.rectangle([x, y, x+w, y+h], outline=(0,0,0), width=3)
        # Draw text
        d.text((x + 20, y + h//2 - 5), text, fill=(0, 0, 0))
        
    img.save(filename)

# Define simple representations
arch_boxes = [
    (100, 100, 200, 100, "Flutter Mobile Client"),
    (400, 100, 200, 100, "Flask REST Backend"),
    (250, 300, 300, 100, "Intelligence Layer\n(ResNet, KNN, LSTM, LLM)"),
    (400, 450, 200, 100, "SQLite Database")
]

er_boxes = [
    (300, 100, 150, 80, "USER"),
    (100, 300, 150, 80, "LISTING"),
    (300, 300, 150, 80, "PURCHASE"),
    (500, 300, 150, 80, "MESSAGE"),
    (300, 500, 150, 80, "TRANSACTION")
]

dfd_boxes = [
    (100, 250, 150, 80, "Farmer / Buyer"),
    (400, 250, 200, 100, "KrishiMitra System"),
    (700, 150, 100, 80, "Groq API"),
    (700, 350, 100, 80, "Agmarknet")
]

flow_boxes = [
    (300, 50, 200, 60, "1. Capture Image"),
    (300, 150, 200, 60, "2. API Transmission"),
    (300, 250, 200, 60, "3. ResNet Inference"),
    (300, 350, 200, 60, "4. Grad-CAM Overlay"),
    (300, 450, 200, 60, "5. JSON Response")
]

create_diagram("basic_arch.png", "System Architecture (Block Diagram)", arch_boxes)
create_diagram("basic_er.png", "Entity-Relationship Model", er_boxes)
create_diagram("basic_dfd.png", "Data Flow Diagram Level 0", dfd_boxes)
create_diagram("basic_flow.png", "Disease Pipeline Flowchart", flow_boxes)

diagrams = [
    ("basic_arch.png", "System Architecture Design"),
    ("basic_er.png", "Entity-Relationship (ER) Diagram"),
    ("basic_dfd.png", "Data Flow Diagram (Level 0 Context Diagram)"),
    ("basic_flow.png", "Flowchart: Disease Detection Pipeline")
]

# Append to document
doc_path = "KrishiMitra_Final_Report_Massive_With_Figures.docx"
print(f"Opening {doc_path}...")
try:
    doc = Document(doc_path)
except Exception as e:
    print(f"Error opening doc: {e}")
    exit(1)

doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Appendix D: System Diagrams and Architecture")
r.bold = True
r.font.size = Pt(14)
r.font.name = "Times New Roman"

p2 = doc.add_paragraph()
r2 = p2.add_run("The following basic architectural diagrams, ER models, DFDs, and flowcharts outline the structural logic and data pathways of the KrishiMitra platform.")
r2.font.size = Pt(12)
r2.font.name = "Times New Roman"
p2.paragraph_format.line_spacing = Pt(21)

fig_idx = 1
for filename, caption in diagrams:
    if os.path.exists(filename):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filename, width=Inches(5.0))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(f"Figure D.{fig_idx}: {caption}")
        r_cap.font.size = Pt(12)
        r_cap.font.name = "Times New Roman"
        r_cap.bold = True
        fig_idx += 1
        doc.add_paragraph()

out_name = "KrishiMitra_Final_Report_Massive_FINAL.docx"
doc.save(out_name)
print(f"Successfully generated simple diagrams and saved as {out_name}!")
