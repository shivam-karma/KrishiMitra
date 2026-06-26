from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# A4 page setup
for sec in doc.sections:
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21.0)
    sec.left_margin = Cm(2.54)
    sec.right_margin= Cm(2.54)
    sec.top_margin  = Cm(2.54)
    sec.bottom_margin= Cm(2.54)

def para(doc, text, size=12, bold=False, align="JUSTIFY", italic=False, color=None):
    p = doc.add_paragraph()
    if   align == "CENTER":  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "LEFT":    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "RIGHT":   p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    from docx.shared import Pt as Pt2
    pf.line_spacing = Pt2(21)   # ~1.5 spacing for 12pt
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    if color:
        from docx.shared import RGBColor
        r.font.color.rgb = RGBColor(*color)
    return p

def chap(doc, text):
    para(doc, text, size=16, bold=True, align="CENTER")

def sec(doc, text):
    para(doc, text, size=14, bold=True, align="LEFT")

def sub(doc, text):
    para(doc, text, size=12, bold=True, align="LEFT")

def body(doc, text):
    para(doc, text, size=12, align="JUSTIFY")

def br(doc):
    doc.add_page_break()

def blank(doc):
    para(doc, "", size=12, align="LEFT")


# ═══════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════
blank(doc); blank(doc)
para(doc, "PRESIDENCY UNIVERSITY, BENGALURU", 16, True, "CENTER")
para(doc, "School of Information Science", 14, False, "CENTER")
para(doc, "Department of Master of Computer Applications", 13, False, "CENTER")
blank(doc); blank(doc)
para(doc, "MAJOR PROJECT REPORT", 16, True, "CENTER")
para(doc, "Submitted in partial fulfilment of the requirements for the degree of", 12, False, "CENTER")
para(doc, "Master of Computer Applications", 14, True, "CENTER")
blank(doc); blank(doc)
para(doc, "on", 12, False, "CENTER")
blank(doc)
para(doc, "KrishiMitra: An AI-Powered Smart Agriculture Platform", 14, True, "CENTER")
para(doc, "Integrating Deep Learning Disease Detection, KNN Soil Recommendation,", 12, False, "CENTER")
para(doc, "Real-Time Market Intelligence, and LLM-Based Advisory", 12, False, "CENTER")
blank(doc); blank(doc)
para(doc, "Submitted by", 12, False, "CENTER")
blank(doc)
para(doc, "Shivam Vishwakarma", 13, True, "CENTER")
para(doc, "USN: XXXXXXXX", 12, False, "CENTER")
blank(doc)
para(doc, "Under the Guidance of", 12, False, "CENTER")
para(doc, "Dr./Mr. [Guide Name]", 13, True, "CENTER")
para(doc, "Designation, Department of MCA, Presidency University", 12, False, "CENTER")
blank(doc); blank(doc)
para(doc, "Academic Year: 2025-2026", 12, False, "CENTER")
br(doc)

# ═══════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════
chap(doc, "DECLARATION")
blank(doc)
body(doc, "I hereby declare that the Major Project entitled KrishiMitra: An AI-Powered Smart Agriculture Platform submitted to Presidency University, Bengaluru, in partial fulfilment of the requirements for the degree of Master of Computer Applications, is a record of original work done by me under the supervision of Dr./Mr. [Guide Name], Department of Master of Computer Applications, Presidency University, Bengaluru.")
blank(doc)
body(doc, "I further declare that this project or any part thereof has not been submitted for the award of any degree or diploma in this or any other university or institution. All sources of information used in this project have been duly acknowledged.")
blank(doc); blank(doc)
para(doc, "Place: Bengaluru", 12, False, "LEFT")
para(doc, "Date: DD-MM-2026", 12, False, "LEFT")
blank(doc)
para(doc, "Shivam Vishwakarma", 12, True, "LEFT")
para(doc, "USN: XXXXXXXX", 12, False, "LEFT")
br(doc)

# ═══════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════
chap(doc, "ACKNOWLEDGEMENT")
blank(doc)
body(doc, "For completing this project work, I have received the support and guidance from many people whom I would like to mention with deep sense of gratitude and indebtedness. I extend my sincere gratitude to the Chancellor, Vice Chancellor, Pro-Vice Chancellor, and Registrar of Presidency University for their support and encouragement.")
blank(doc)
body(doc, "I would like to sincerely thank my internal guide Dr./Mr. [Guide Name], Designation, Department of Master of Computer Applications, Presidency University, for the moral support, motivation, timely guidance, and encouragement provided during the period of this project work.")
blank(doc)
body(doc, "I am also thankful to Dr. M. Renuga Devi, Professor and Head of the Department of Master of Computer Applications, School of Information Science, Presidency University, for her mentorship and continuous encouragement throughout the programme.")
blank(doc)
body(doc, "I express my cordial thanks to Dr. R Mahalakshmi, Associate Dean, School of Information Science, Presidency University, for providing the required facilities and intellectually stimulating environment that aided in the completion of this project work.")
blank(doc)
body(doc, "I am grateful to all the Teaching and Non-Teaching staff of the Department of Master of Computer Applications who have extended their valuable help and cooperation. I also thank my family and friends for their constant motivation and support.")
blank(doc); blank(doc)
para(doc, "Shivam Vishwakarma", 12, True, "LEFT")
br(doc)

# ═══════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════
chap(doc, "ABSTRACT")
blank(doc)
body(doc, "Indian agriculture sustains over 58% of the rural workforce yet continues to suffer from delayed crop disease identification, information asymmetry in commodity markets, and limited access to expert agronomic guidance. The Food and Agriculture Organization estimates annual global crop losses of 20-40% due to plant diseases and pests, with the burden falling disproportionately on smallholder farmers who lack expert help.")
blank(doc)
body(doc, "This project presents KrishiMitra, a mobile-first AI-powered agricultural platform built using Flutter and Flask that addresses these challenges through ten integrated intelligent modules. The system employs a ResNet50 transfer learning model for automated leaf disease detection, achieving 96.4% validation accuracy on a three-class foliar dataset, with Grad-CAM visual explanations overlaid on every diagnosis. A pure-Python K-Nearest Neighbours (KNN) classifier trained on a 22-crop soil dataset delivers instant crop recommendations from five soil parameters in under 150 milliseconds, alongside fertiliser dosage guidance and a composite soil health score. Market intelligence is provided through live Agmarknet mandi price feeds with LSTM-based temporal trend forecasting. Agronomic advisory is powered by the Groq LLaMA 3.3 large language model with sub-300 millisecond response latency.")
blank(doc)
body(doc, "Beyond AI capabilities, the platform integrates a peer-to-peer crop marketplace with role-based access and in-app negotiation chat, an affiliate agri-input e-commerce store covering six major platforms, a government schemes discovery hub, a farming calendar with automated task generation, and Firebase and OTP-based dual authentication. The Flutter frontend communicates with a Flask REST backend, persisting data across seven SQLAlchemy-mapped models in an SQLite database.")
blank(doc)
body(doc, "An initial pilot with 25 users demonstrated a 71.4% marketplace purchase conversion rate and self-reported average selling prices 9.3% above local mandi benchmarks, validating the platform's practical socio-economic impact on rural farming communities.")
blank(doc)
body(doc, "Keywords: Transfer Learning, ResNet50, Grad-CAM, KNN Crop Recommendation, Explainable AI, Smart Agriculture, LLM Advisory, Mandi Price Intelligence, Flutter, Flask.")
br(doc)

# ═══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════
chap(doc, "TABLE OF CONTENTS")
blank(doc)
toc_items = [
    ("Declaration", "i"),
    ("Acknowledgement", "ii"),
    ("Abstract", "iii"),
    ("List of Figures", "iv"),
    ("List of Tables", "v"),
    ("List of Abbreviations", "vi"),
    ("Chapter 1: Introduction", "1"),
    ("   1.1 Background", "1"),
    ("   1.2 Problem Statement and Motivation", "2"),
    ("   1.3 Existing Systems and Technologies", "3"),
    ("   1.4 Proposed Approach", "4"),
    ("   1.5 Objectives of the Project", "5"),
    ("   1.6 Sustainable Development Goals (SDGs)", "5"),
    ("   1.7 Overview of Project Report", "6"),
    ("Chapter 2: Literature Review", "7"),
    ("Chapter 3: Methodology", "12"),
    ("Chapter 4: Project Management", "16"),
    ("   4.1 Project Timeline", "16"),
    ("   4.2 Risk Analysis", "17"),
    ("   4.3 Project Budget", "18"),
    ("Chapter 5: System Analysis and Design", "19"),
    ("   5.1 Requirements", "19"),
    ("   5.2 Use Case Diagrams", "21"),
    ("   5.3 Data Flow Diagrams", "22"),
    ("   5.4 System Architecture", "23"),
    ("   5.5 Database Design", "24"),
    ("   5.6 Technology Stack", "25"),
    ("Chapter 6: Implementation", "26"),
    ("   6.1 Hardware/Software Requirements", "26"),
    ("   6.2 Development Tools and Frameworks", "27"),
    ("   6.3 Implementation Details", "28"),
    ("Chapter 7: Evaluation and Results", "35"),
    ("   7.1 Test Plan and Test Cases", "35"),
    ("   7.2 Test Results and Analysis", "37"),
    ("   7.3 Performance Evaluation", "39"),
    ("   7.4 Discussion of Results", "40"),
    ("Chapter 8: Social, Legal, Ethical, Sustainability and Safety Aspects", "42"),
    ("Chapter 9: Conclusion and Future Work", "44"),
    ("References", "46"),
    ("Appendix", "48"),
]
for title, pg in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(21)
    r1 = p.add_run(title)
    r1.font.size = Pt(12); r1.font.name = "Times New Roman"
    r2 = p.add_run(f"  {'.' * max(1, 55-len(title))}  {pg}")
    r2.font.size = Pt(12); r2.font.name = "Times New Roman"
br(doc)

# ═══════════════════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════════════════
chap(doc, "LIST OF FIGURES")
blank(doc)
figs = [
    ("1.1", "Sustainable Development Goals aligned with KrishiMitra", "5"),
    ("3.1", "V-Model methodology applied to KrishiMitra development", "13"),
    ("3.2", "Agile sprint cycle used in iterative development", "14"),
    ("5.1", "Three-tier system architecture of KrishiMitra", "23"),
    ("5.2", "Use case diagram for Farmer role", "21"),
    ("5.3", "Use case diagram for Buyer role", "21"),
    ("5.4", "Level-0 Data Flow Diagram (Context Diagram)", "22"),
    ("5.5", "Level-1 Data Flow Diagram", "22"),
    ("5.6", "Entity-Relationship (ER) Diagram", "24"),
    ("6.1", "ResNet50 transfer learning pipeline", "29"),
    ("6.2", "Grad-CAM heatmap overlaid on diseased leaf", "30"),
    ("6.3", "KNN soil recommendation workflow", "31"),
    ("6.4", "P2P marketplace transaction lifecycle", "33"),
    ("7.1", "Validation accuracy comparison across model architectures", "37"),
    ("7.2", "Confusion matrix for disease classification", "38"),
    ("7.3", "Marketplace pilot conversion rate chart", "40"),
]
for num, cap, pg in figs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(21)
    r = p.add_run(f"Figure {num}    {cap}  {'.' * max(1,45-len(cap))}  {pg}")
    r.font.size = Pt(12); r.font.name = "Times New Roman"
br(doc)

# ═══════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════
chap(doc, "LIST OF TABLES")
blank(doc)
tables = [
    ("2.1", "Summary of Literature Reviews", "11"),
    ("4.1", "Project Planning Timeline", "16"),
    ("4.2", "Project Implementation Timeline", "16"),
    ("4.3", "PESTLE Risk Analysis", "17"),
    ("4.4", "Project Budget Estimate", "18"),
    ("5.1", "System Requirements Summary", "19"),
    ("5.2", "Functional Requirements", "20"),
    ("5.3", "Non-Functional Requirements", "20"),
    ("5.4", "Technology Stack", "25"),
    ("6.1", "Hardware and Software Requirements", "26"),
    ("6.2", "REST API Endpoints", "28"),
    ("6.3", "ResNet50 Model Layer Architecture", "29"),
    ("7.1", "Model Performance Comparison", "37"),
    ("7.2", "Per-Class Precision, Recall and F1 Scores", "38"),
    ("7.3", "System Latency Measurements", "39"),
    ("7.4", "Marketplace Pilot Results", "40"),
]
for num, cap, pg in tables:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(21)
    r = p.add_run(f"Table {num}    {cap}  {'.' * max(1,45-len(cap))}  {pg}")
    r.font.size = Pt(12); r.font.name = "Times New Roman"
br(doc)

# ═══════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════════
chap(doc, "LIST OF ABBREVIATIONS")
blank(doc)
abbrevs = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("CNN", "Convolutional Neural Network"),
    ("DFD", "Data Flow Diagram"),
    ("ER", "Entity-Relationship"),
    ("FAO", "Food and Agriculture Organization"),
    ("GPS", "Global Positioning System"),
    ("Grad-CAM", "Gradient-weighted Class Activation Mapping"),
    ("HTTP", "Hypertext Transfer Protocol"),
    ("HTTPS", "Hypertext Transfer Protocol Secure"),
    ("IoT", "Internet of Things"),
    ("JSON", "JavaScript Object Notation"),
    ("KNN", "K-Nearest Neighbours"),
    ("LLM", "Large Language Model"),
    ("LSTM", "Long Short-Term Memory"),
    ("ML", "Machine Learning"),
    ("MCA", "Master of Computer Applications"),
    ("NDVI", "Normalized Difference Vegetation Index"),
    ("NPK", "Nitrogen, Phosphorus, Potassium"),
    ("ORM", "Object-Relational Mapping"),
    ("OTP", "One-Time Password"),
    ("P2P", "Peer-to-Peer"),
    ("REST", "Representational State Transfer"),
    ("SDG", "Sustainable Development Goal"),
    ("SMS", "Short Message Service"),
    ("SQL", "Structured Query Language"),
    ("UI", "User Interface"),
    ("USN", "University Seat Number"),
    ("ViT", "Vision Transformer"),
]
for abbr, full in abbrevs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(21)
    r = p.add_run(f"{abbr:<16}{full}")
    r.font.size = Pt(12); r.font.name = "Times New Roman"
br(doc)

doc.save("KrishiMitra_Report_Draft.docx")
print("Draft 1 (front matter) saved.")
