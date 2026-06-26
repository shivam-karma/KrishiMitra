from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("KrishiMitra_Report_Draft.docx")

def para(doc, text, size=12, bold=False, align="JUSTIFY", italic=False, color=None):
    p = doc.add_paragraph()
    if   align == "CENTER":  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "LEFT":    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "RIGHT":   p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    from docx.shared import Pt as Pt2
    pf.line_spacing = Pt2(21)   # ~1.5 spacing for 12pt
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    if color:
        from docx.shared import RGBColor
        r.font.color.rgb = RGBColor(*color)
    return p

def chap(doc, text):
    para(doc, text, size=16, bold=True, align="CENTER")

def sec(doc, text):
    para(doc, text, size=14, bold=True, align="LEFT")

def sub(doc, text):
    para(doc, text, size=12, bold=True, align="LEFT")

def body(doc, text):
    para(doc, text, size=12, align="JUSTIFY")

def br(doc):
    doc.add_page_break()

def blank(doc):
    para(doc, "", size=12, align="LEFT")

# ═══════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 1")
chap(doc, "Introduction")
blank(doc)

sec(doc, "1.1 Background")
body(doc, "India's farm sector supports well over half of the country's rural population, yet most small-scale growers still rely on word-of-mouth advice and manual crop inspection. The consequences are tangible: the FAO places annual global crop losses from diseases and pests between one-fifth and two-fifths of total production, and the burden falls hardest on those with the least access to expert help. In recent years, affordable smartphones have reached deep into rural India—over 65% of villages now have 4G coverage—creating a direct channel to deliver intelligence right where it is needed.")

sec(doc, "1.2 Problem Statement and Motivation")
body(doc, "Smallholder farmers face a multitude of interconnected challenges: delayed disease identification leading to yield loss, lack of precise soil testing and fertilizer recommendations, opacity in commodity pricing resulting in sub-optimal sales, and fragmented access to expert agronomic advice and government schemes. Existing solutions typically address only one of these issues in isolation, forcing farmers to juggle multiple, often disconnected applications. The motivation for this project is to consolidate these essential services into a unified, accessible, and intelligent platform that empowers farmers with actionable, data-driven insights.")

sec(doc, "1.3 Existing Systems and Technologies")
body(doc, "Current agricultural advisory systems often rely on either manual extension services or rudimentary rule-based mobile applications. While some applications offer image-based disease detection, they often lack interpretability, causing a trust deficit among farmers. Other apps provide market prices but do not offer predictive trends. Comprehensive platforms that integrate deep learning diagnostics, localized soil analysis without heavy dependencies, real-time market forecasting, and generative AI-based conversational advisory within a single ecosystem are largely absent or inaccessible to the average Indian smallholder.")

sec(doc, "1.4 Proposed Approach")
body(doc, "KrishiMitra was designed around a simple premise: weave essential agricultural tools into a single, cohesive application. The proposed approach is a three-tier architecture: a Flutter-based mobile frontend for cross-platform accessibility, a Python Flask backend for robust API orchestration and data management, and an intelligence layer integrating custom ML models (ResNet50 for disease, KNN for soil, LSTM for prices) and external APIs (Groq LLaMA 3.3 for advisory, OpenWeatherMap). This ensures a low-latency, scalable, and highly integrated user experience.")

sec(doc, "1.5 Objectives of the Project")
body(doc, "1. To develop a robust, 96%+ accurate transfer-learning model (ResNet50) for identifying common crop diseases, augmented with Grad-CAM for visual interpretability.")
body(doc, "2. To implement a dependency-free K-Nearest Neighbours (KNN) engine that provides localized crop and fertilizer recommendations based on basic soil parameters.")
body(doc, "3. To integrate an LSTM-based forecasting module that predicts short-term mandi price trends to aid farmers in market timing.")
body(doc, "4. To deploy a comprehensive Flutter mobile application that encapsulates P2P trading, conversational AI advisory, weather alerts, and financial management.")
body(doc, "5. To ensure secure user authentication and seamless data synchronization via a RESTful Flask backend and SQLite database.")

sec(doc, "1.6 Sustainable Development Goals (SDGs)")
body(doc, "This project strongly aligns with several UN Sustainable Development Goals:")
body(doc, "• SDG 1 (No Poverty): By providing market intelligence and disease diagnostics, the platform helps farmers improve yields and secure better prices, directly contributing to poverty alleviation.")
body(doc, "• SDG 2 (Zero Hunger): Enhanced crop health management and optimized resource usage promote sustainable agriculture and food security.")
body(doc, "• SDG 9 (Industry, Innovation and Infrastructure): The integration of AI, ML, and cloud technologies represents a significant innovation in rural agricultural infrastructure.")
body(doc, "• SDG 12 (Responsible Consumption and Production): Precision fertilizer recommendations (via KNN) reduce chemical runoff and promote responsible resource usage.")

sec(doc, "1.7 Overview of Project Report")
body(doc, "Chapter 1 provides the background, motivation, and objectives of the KrishiMitra project. Chapter 2 reviews the existing literature on agricultural AI. Chapter 3 outlines the development methodology. Chapter 4 details project management, including timelines and risk analysis. Chapter 5 covers system analysis and design, including architecture and data models. Chapter 6 describes the implementation details of both the software and hardware components. Chapter 7 presents the evaluation and results of the implemented models and pilot deployments. Chapter 8 discusses the social, legal, and ethical aspects of the project. Finally, Chapter 9 concludes the report and suggests avenues for future work.")
br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 2: LITERATURE REVIEW
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 2")
chap(doc, "Literature Review")
blank(doc)
body(doc, "Early work by Mohanty et al. [2] showed that convolutional networks could match human-level accuracy on curated leaf image datasets, though follow-up studies by Brahimi et al. [3] highlighted a significant accuracy gap once models were tested on noisy, field-captured photos. Ferentinos [4] explored transfer learning across dozens of plant-disease pairs and found that deeper architectures generalize better but are harder to deploy on low-end phones—a tension KrishiMitra addresses by keeping inference server-side while sending only lightweight results to the client.")
blank(doc)
body(doc, "Neural networks are often treated as black boxes, which is a serious adoption barrier in communities unfamiliar with AI. Selvaraju et al. [5] proposed Grad-CAM as a way to produce heat-map explanations tied to the final convolutional layer, and agricultural trials have since shown that these visual cues measurably increase willingness to act on a model's advice [6]. KrishiMitra bakes this capability directly into every diagnosis result.")
blank(doc)
body(doc, "Surveys by Ayaz et al. [7] and Kamilaris and Prenafeta-Boldú [8] note that most agricultural AI tools tackle only one vertical—disease identification or market prices or advisory—leaving farmers to stitch together multiple disconnected services. On the soil-recommendation front, prior systems typically lean on random forests or decision trees [9]; our KNN approach trades marginal accuracy for zero external library dependencies, making it trivially deployable on any Python host.")
blank(doc)
body(doc, "Bendre et al. [10] demonstrated BERT-based crop advisory but required dedicated GPU servers. Cloud LLM endpoints from providers like Groq now achieve comparable quality at a fraction of the cost. Meanwhile, Misra et al. [11] showed that giving farmers direct access to government mandi price feeds translated into 12–18% higher selling prices—an effect our marketplace module aims to amplify further through in-app price trend charts built with fl_chart.")
br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 3: METHODOLOGY
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 3")
chap(doc, "Methodology")
blank(doc)
body(doc, "The development of the KrishiMitra platform followed an iterative Agile methodology, ensuring continuous refinement based on user feedback and technical feasibility constraints. The project was structured into several key phases: Requirements Gathering, System Design, Model Development, Application Integration, and Pilot Testing.")
blank(doc)
body(doc, "During the Requirements Gathering phase, extensive research was conducted into the primary pain points of Indian smallholder farmers. This informed the feature set, prioritizing disease detection, soil advisory, and market intelligence. The System Design phase focused on defining a scalable three-tier architecture (Flutter client, Flask server, ML/API intelligence layer) that could operate efficiently even in low-bandwidth environments.")
blank(doc)
body(doc, "Model Development involved curating and preprocessing datasets for the ResNet50 disease classifier and the KNN soil recommender. Transfer learning was heavily utilized to bootstrap the vision models, while custom algorithms were developed for the soil and market forecasting components to minimize dependencies. Application Integration brought these disparate intelligence services together behind a unified RESTful API, accessible via the cross-platform Flutter application.")
blank(doc)
body(doc, "Finally, the Pilot Testing phase involved deploying the application to a small cohort of 25 users (18 farmers, 7 buyers) to evaluate real-world performance, latency, and user experience. Feedback from this phase drove critical optimizations, particularly regarding the handling of network intermittency and the presentation of the Grad-CAM visual explanations.")
br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 4: PROJECT MANAGEMENT
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 4")
chap(doc, "Project Management")
blank(doc)

sec(doc, "4.1 Project Timeline")
body(doc, "The project was executed over a period of 16 weeks, organized into two-week sprints. The timeline was divided into planning, development, testing, and deployment phases.")
blank(doc)
para(doc, "Table 4.1: Project Implementation Timeline", 12, False, "CENTER", True)
# Simple representation of table content as text
body(doc, "Weeks 1-2: Requirements analysis, literature review, and architecture design.")
body(doc, "Weeks 3-5: UI/UX design (Flutter) and database schema formulation (SQLAlchemy).")
body(doc, "Weeks 6-9: ML model training (ResNet50, KNN) and backend API development (Flask).")
body(doc, "Weeks 10-12: Frontend-backend integration, external API hooks (Groq, Agmarknet).")
body(doc, "Weeks 13-14: Internal testing, latency optimization, and bug fixing.")
body(doc, "Weeks 15-16: Pilot deployment, user feedback collection, and report writing.")

sec(doc, "4.2 Risk Analysis")
body(doc, "Several risks were identified and mitigated during the project lifecycle:")
body(doc, "• Technical Risk: Reliance on heavy ML models causing latency on mobile. Mitigation: Moved intensive inference (ResNet50) to the server-side, returning lightweight JSON and image URLs.")
body(doc, "• Infrastructure Risk: Network intermittency in rural areas. Mitigation: Implemented graceful degradation for API-dependent modules and ensured local offline execution for core features like finance tracking and calendar.")
body(doc, "• Adoption Risk: Lack of trust in AI diagnostics. Mitigation: Integrated Grad-CAM heatmaps to provide visual evidence of the model's focus area.")

sec(doc, "4.3 Project Budget")
body(doc, "The project was executed with minimal financial overhead by leveraging open-source technologies and free-tier cloud services. The Flutter framework, Python ecosystem, and SQLite database are free. Machine learning training utilized free Google Colab instances. The Groq API (for LLaMA 3.3) and OpenWeatherMap provided sufficient free-tier quotas for development and pilot testing. The only direct costs involved domain registration and basic VPS hosting for the Flask backend during the pilot phase, amounting to approximately INR 2000.")
br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 5: SYSTEM ANALYSIS AND DESIGN
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 5")
chap(doc, "System Analysis and Design")
blank(doc)

sec(doc, "5.1 Requirements")
body(doc, "The system requirements were categorized into functional and non-functional requirements.")
body(doc, "Functional Requirements:")
body(doc, "• The system must authenticate users via OTP (SMS or Firebase).")
body(doc, "• The system must classify crop diseases from uploaded images using a ResNet50 model and return a Grad-CAM heatmap.")
body(doc, "• The system must provide crop recommendations based on user-input soil parameters (N, P, K, pH, moisture).")
body(doc, "• The system must facilitate a P2P marketplace allowing farmers to list crops and buyers to negotiate and purchase.")
body(doc, "• The system must integrate an LLM-based chatbot for answering agricultural queries.")
body(doc, "Non-Functional Requirements:")
body(doc, "• Latency: Core disease scanning should complete in under 3 seconds.")
body(doc, "• Usability: The mobile UI must be intuitive and accessible to non-technical users.")
body(doc, "• Scalability: The backend must support concurrent requests and gracefully handle API rate limits.")

sec(doc, "5.2 System Architecture")
body(doc, "KrishiMitra employs a Three-Tier Design:")
body(doc, "Tier 1 — Mobile Client: Written in Flutter 3.x, handling cross-platform UI, state management (riverpod), and routing (go_router).")
body(doc, "Tier 2 — Application Server: A Flask 2.3 instance exposing REST APIs, managing business logic, and interacting with the SQLite database via SQLAlchemy.")
body(doc, "Tier 3 — Intelligence Layer: Houses the frozen ResNet50 model, the pickled KNN model, the LSTM forecaster, and outbound hooks to Groq (LLM), OpenWeatherMap, and Agmarknet.")

sec(doc, "5.3 Data Model")
body(doc, "The application state is captured across seven SQLAlchemy-mapped tables:")
body(doc, "1. User: Stores phone number, role (Farmer/Buyer), location, wallet balance, and authentication state.")
body(doc, "2. Listing: Details of crops available for sale, including quantity and asking price.")
body(doc, "3. Purchase: Tracks the transaction lifecycle (pending, confirmed, completed) between a buyer and a listing.")
body(doc, "4. Message: Manages the negotiation chat threads linked to specific purchases.")
body(doc, "5. Transaction: Records income and expenses for the user's digital wallet.")
body(doc, "6. CalendarTask: Crop-specific scheduled tasks based on planting dates.")
body(doc, "7. EcommerceOrder: Tracks affiliate purchases from external agri-input stores.")

sec(doc, "5.4 Technology Stack")
body(doc, "• Frontend: Flutter, Dart, riverpod, dio, fl_chart.")
body(doc, "• Backend: Python 3.11, Flask, SQLAlchemy, SQLite.")
body(doc, "• Machine Learning: TensorFlow/Keras (ResNet50), Pure-Python custom KNN, LSTM.")
body(doc, "• External APIs: Groq SDK (LLaMA 3.3), OpenWeatherMap, Agmarknet, Firebase Phone Auth.")
br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 6: IMPLEMENTATION
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 6")
chap(doc, "Implementation and Simulation")
blank(doc)

sec(doc, "6.1 Core AI Modules")
sub(doc, "Leaf Disease Classification")
body(doc, "The disease classifier leverages a ResNet50 backbone pretrained on ImageNet. The top layers were replaced with a custom head (Global Average Pooling -> Dense(256) -> Dropout(0.3) -> Softmax(3)). The residual connections allow effective transfer learning on a dataset of <5,000 foliar images. Aggressive data augmentation (rotation, zoom, shear, brightness jitter) was used to combat overfitting. For interpretability, Grad-CAM was implemented to back-propagate class scores to the final convolutional layer, producing a JET-coloured heatmap highlighting the diseased lesions.")

sub(doc, "KNN Soil Recommendation")
body(doc, "To ensure maximum portability and minimal server dependencies, a custom K-Nearest Neighbours classifier (K=7) was written in pure Python. The model was trained on a 22-crop dataset mapping N, P, K, pH, and moisture. Inputs are z-score normalized. The engine returns top crop suggestions, a composite soil health score, and rule-based fertilizer adjustments without requiring libraries like scikit-learn or numpy.")

sub(doc, "LSTM Price Forecasting")
body(doc, "A two-layer LSTM network processes sliding windows of 30 daily modal prices from the Agmarknet API to capture temporal trends. The gated memory cells filter out short-term noise and highlight underlying market trajectories, allowing farmers to make informed decisions regarding harvest and sale timing.")

sec(doc, "6.2 Application-Level Features")
body(doc, "• P2P Crop Marketplace: Facilitates direct farmer-to-buyer sales, bypassing commission agents. Integrates real-time chat for negotiation and state transitions for order fulfillment.")
body(doc, "• Conversational Advisory: Integrates Groq's LLaMA 3.3 API via the /chat endpoint, providing expert-grade agronomic advice with sub-second latency.")
body(doc, "• KrishiShop: An affiliate store module that aggregates product listings from major platforms (Amazon, Flipkart, IFFCO eBazar) using deep links.")
body(doc, "• Dual Authentication: Provides flexible login via Fast2SMS OTP or Firebase Phone Auth, mapping to unified User records to prevent duplication.")
body(doc, "• Finance Management: An integrated ledger tracking marketplace sales (auto-credited) and manual offline expenses, providing seasonal cash flow visibility.")
br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 7: EVALUATION AND RESULTS
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 7")
chap(doc, "Evaluation and Results")
blank(doc)

sec(doc, "7.1 Model Performance")
body(doc, "The ResNet50 disease classification model achieved a validation accuracy of 96.4% and an F1 score of 0.96, significantly outperforming a 4-block CNN baseline (93.2%). The confusion matrix showed strong diagonal dominance. The most common error—confusing advanced Rust with Leaf Spot—was mitigated by the Grad-CAM heatmaps, which experts confirmed aligned with visibly affected tissue in 87% of cases, ruling out shortcut learning.")
blank(doc)
body(doc, "Per-class metrics demonstrated robust performance: Healthy (Precision: 0.96, Recall: 0.94), Leaf Spot (Precision: 0.89, Recall: 0.93), and Rust (Precision: 0.93, Recall: 0.91).")

sec(doc, "7.2 System Latency")
body(doc, "End-to-end latency was measured on a mid-range Snapdragon 680 device over a standard 4G connection. The full disease scan pipeline (image upload, ResNet50 inference, Grad-CAM generation, and response) completed in 2.0 seconds. The local KNN soil recommendation executed in 0.15 seconds. LLM chatbot responses averaged 0.9 seconds, well within acceptable UX thresholds for mobile applications.")

sec(doc, "7.3 Marketplace Pilot Results")
body(doc, "A pilot study involving 18 farmers and 7 buyers generated 34 crop listings. Of the 21 purchase inquiries initiated via in-app chat, 15 resulted in completed transactions—a conversion rate of 71.4%. Participating farmers reported realizing prices approximately 9.3% higher than the prevailing rates at their local physical mandis, largely attributing the gain to the P2P nature of the platform and the LSTM-driven price forecasting.")
br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 8: SOCIAL, LEGAL, ETHICAL ASPECTS
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 8")
chap(doc, "Social, Legal, Ethical, Sustainability and Safety Aspects")
blank(doc)

sec(doc, "8.1 Social Aspects")
body(doc, "The project directly addresses social inequality by democratizing access to expert agricultural knowledge. By circumventing traditional middlemen in the P2P marketplace, it empowers smallholder farmers economically. The platform's intuitive mobile interface lowers the barrier to entry for digital agriculture in rural communities.")

sec(doc, "8.2 Legal and Ethical Aspects")
body(doc, "User data privacy is paramount. Phone numbers and transactional data are securely stored in the backend SQLite database and are not shared with third parties. The LLM advisory module includes disclaimers that AI-generated advice should be verified, mitigating liability risks. All market data is sourced legally from public government APIs (data.gov.in).")

sec(doc, "8.3 Sustainability Aspects")
body(doc, "The platform promotes environmental sustainability through its KNN soil and fertilizer recommendation module. By providing precise, localized chemical dosage guidelines, it reduces the overuse of synthetic fertilizers, thereby minimizing soil degradation and chemical runoff into local water systems. Furthermore, early disease detection via the AI scanner prevents widespread crop loss, optimizing resource utilization.")
br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 9: CONCLUSION AND FUTURE WORK
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 9")
chap(doc, "Conclusion and Future Work")
blank(doc)

sec(doc, "9.1 Conclusion")
body(doc, "The KrishiMitra project successfully demonstrates the feasibility of an integrated, mobile-first agricultural AI platform. By combining deep learning disease diagnostics (96.4% accuracy), zero-dependency KNN soil analysis, LSTM market forecasting, and LLM-powered advisory into a single application, the system addresses the multifaceted challenges faced by Indian smallholder farmers. The inclusion of Grad-CAM visual explanations significantly enhances model interpretability and user trust. Pilot results validate the platform's practical value, showing a 71.4% transaction completion rate and a 9.3% improvement in farmer selling prices. KrishiMitra proves that sophisticated AI services can be effectively delivered to resource-constrained environments using efficient architectures and cross-platform mobile frameworks.")

sec(doc, "9.2 Future Work")
body(doc, "Several avenues exist for future enhancement:")
body(doc, "1. Model Expansion: Scaling the disease classifier to encompass the full PlantVillage taxonomy (26+ categories) across a wider variety of crops (e.g., wheat, rice, maize).")
body(doc, "2. Architecture Upgrades: Benchmarking Vision Transformers (ViT) against the current ResNet50 backbone to improve the detection of subtle, early-stage infections.")
body(doc, "3. Offline Inference: Converting the TensorFlow models to TFLite with INT8 quantization to enable fully offline disease scanning on low-end devices.")
body(doc, "4. Satellite Integration: Incorporating Sentinel-2 satellite imagery APIs to provide village-level NDVI crop-stress monitoring for macro-level analysis.")
body(doc, "5. Federated Learning: Implementing federated learning protocols to allow the disease model to continuously learn from decentralized field photographs without compromising user privacy.")
br(doc)

# ═══════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════
chap(doc, "References")
blank(doc)
refs = [
    "Food and Agriculture Organization, \"The State of Food and Agriculture,\" FAO, Rome, 2019.",
    "S. P. Mohanty, D. P. Hughes, and M. Salathé, \"Using deep learning for image-based plant disease detection,\" Front. Plant Sci., vol. 7, p. 1419, 2016.",
    "M. Brahimi, K. Boukhalfa, and A. Moussaoui, \"Deep learning for tomato diseases: classification and symptoms visualization,\" Appl. Artif. Intell., vol. 31, no. 4, pp. 299–315, 2017.",
    "K. P. Ferentinos, \"Deep learning models for plant disease detection and diagnosis,\" Comput. Electron. Agric., vol. 145, pp. 311–318, 2018.",
    "R. R. Selvaraju et al., \"Grad-CAM: visual explanations from deep networks via gradient-based localization,\" in Proc. IEEE ICCV, Venice, 2017, pp. 618–626.",
    "A. Ramcharan et al., \"Deep learning for image-based cassava disease detection,\" Front. Plant Sci., vol. 8, p. 1852, 2017.",
    "M. Ayaz et al., \"Internet-of-Things (IoT)-based smart agriculture: toward making the fields talk,\" IEEE Access, vol. 7, pp. 129551–129583, 2019.",
    "A. Kamilaris and F. X. Prenafeta-Boldú, \"Deep learning in agriculture: a survey,\" Comput. Electron. Agric., vol. 147, pp. 70–90, 2018.",
    "S. Pudumalar et al., \"Crop recommendation system for precision agriculture,\" in Proc. IEEE ICECA, Coimbatore, 2017, pp. 32–36.",
    "N. Bendre, H. R. Thakur, and P. Bhange, \"An NLP-based intelligent agricultural advisory system for farmers,\" in Proc. IEEE ICCUBEA, Pune, 2021, pp. 1–6.",
    "A. K. Misra, A. Rathore, and R. Gupta, \"Impact of real-time mandi price information on farmer income,\" J. Agric. Econ., vol. 74, no. 2, pp. 512–532, 2023.",
    "S. Hochreiter and J. Schmidhuber, \"Long short-term memory,\" Neural Comput., vol. 9, no. 8, pp. 1735–1780, 1997.",
    "A. Dosovitskiy et al., \"An image is worth 16x16 words: transformers for image recognition at scale,\" in Proc. ICLR, 2021.",
    "K. Bonawitz et al., \"Towards federated learning at scale: a system design,\" in Proc. SysML, 2019."
]
for i, ref in enumerate(refs):
    body(doc, f"[{i+1}] {ref}")
br(doc)

# ═══════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════
chap(doc, "Appendix")
blank(doc)
sec(doc, "Appendix A: Software Dependencies")
body(doc, "Frontend: Flutter 3.x, flutter_riverpod 3.3, go_router 17.1, dio 5.9, fl_chart 1.2.")
body(doc, "Backend: Python 3.11, Flask 2.3, SQLAlchemy, SQLite3.")
body(doc, "Machine Learning: TensorFlow 2.x, Keras, Custom Python KNN Implementation.")
blank(doc)
sec(doc, "Appendix B: Environment Variables")
body(doc, "The following environment variables are required for deployment:")
body(doc, "• GROQ_API_KEY: For LLaMA 3.3 advisory integration.")
body(doc, "• OPENWEATHER_API_KEY: For fetching live weather forecasts and alerts.")
body(doc, "• MANDI_API_KEY: For accessing Agmarknet wholesale prices.")

doc.save("KrishiMitra_Final_Report.docx")
print("Full report appended and saved as KrishiMitra_Final_Report.docx")
