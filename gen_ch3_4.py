from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def get_doc(filename="KrishiMitra_Ch3_Ch4.docx"):
    doc = Document()
    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width  = Cm(21.0)
        sec.left_margin = Cm(2.54)
        sec.right_margin= Cm(2.54)
        sec.top_margin  = Cm(2.54)
        sec.bottom_margin= Cm(2.54)
    return doc

def para(doc, text, size=12, bold=False, align="JUSTIFY", italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = {"CENTER": WD_ALIGN_PARAGRAPH.CENTER, "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                   "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.line_spacing = Pt(21)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    if color: r.font.color.rgb = RGBColor(*color)
    return p

def chap(doc, text): para(doc, text, size=16, bold=True, align="CENTER")
def sec(doc, text):  para(doc, text, size=14, bold=True, align="LEFT")
def sub(doc, text):  para(doc, text, size=12, bold=True, align="LEFT")
def body(doc, text): para(doc, text, size=12, align="JUSTIFY")
def blank(doc):      para(doc, "", size=12, align="LEFT")
def br(doc):         doc.add_page_break()

def add_table(doc, headers, rows, caption=""):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.bold = True; r.font.size = Pt(11); r.font.name = "Times New Roman"
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(10); run.font.name = "Times New Roman"
        tc = c._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'FFFFFF')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        for p in c.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255,255,255)
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        fill = 'EEF4FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = "Times New Roman"
            tc = c._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill)
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
    blank(doc)

doc = get_doc()

# ═══════════════════════════════════════════════════
# CHAPTER 3: METHODOLOGY
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 3")
chap(doc, "Methodology")
blank(doc)

body(doc, "The development of KrishiMitra followed the Agile Software Development Life Cycle (SDLC) methodology, specifically adopting the Scrum framework. Agile is widely recognised as the most suitable methodology for projects involving rapidly evolving requirements, frequent integration of AI/ML components, and iterative user feedback — all of which characterise the KrishiMitra development process [1]. Unlike rigid linear models such as Waterfall, Agile enables the development team to incrementally build, test, and refine features across short sprints, ensuring that each component of the application — from the CNN disease detection pipeline to the Flutter UI screens — is validated independently before system integration. The Agile methodology's emphasis on working software over comprehensive documentation, customer collaboration over contract negotiation, and responding to change over following a plan aligns directly with the exploratory and data-driven nature of AI-integrated mobile application development [2].")
blank(doc)

sec(doc, "3.1 Agile Scrum Methodology")
body(doc, "The Scrum framework organises development into time-boxed iterations called sprints, each lasting two to three weeks. At the beginning of each sprint, a sprint planning meeting identifies tasks from the product backlog to be completed. Daily standups ensure team alignment, while sprint review and retrospective meetings at the end of each sprint allow for continuous improvement. For KrishiMitra, the product backlog was derived from the five core objectives defined in Chapter 1: AI disease detection, soil crop recommendation, P2P marketplace, OTP authentication, and the Flask REST API backend. Each sprint targeted a specific functional module, ensuring that at the end of every sprint, a demonstrable and tested increment of the software was available [3].")
blank(doc)

para(doc, "Figure 3.1: Agile Scrum Methodology applied to KrishiMitra development", size=11, italic=True, align="CENTER")
blank(doc)

body(doc, "Figure 3.1 illustrates the Agile Scrum cycle as applied to the KrishiMitra project. The cycle begins with the Product Backlog containing all feature requirements, which are then selected for each Sprint Backlog. During the sprint, daily development and testing activities occur, culminating in a Sprint Review where a working increment is demonstrated. Feedback is incorporated into the next sprint, ensuring continuous alignment with project objectives [4].")
blank(doc)

sec(doc, "3.2 Mapping KrishiMitra Stages to Agile Phases")
body(doc, "The following table maps the key stages of KrishiMitra's development to the corresponding phases of the Agile Scrum methodology, demonstrating how each technical component was addressed within the iterative framework.")
blank(doc)

add_table(doc,
    ["Sprint", "Phase", "Activities", "Deliverable"],
    [
        ["Sprint 0", "Requirements & Design", "Gather functional requirements, define user stories, design system architecture, select technology stack (Flutter, Flask, TensorFlow, SQLite)", "System architecture diagram, API endpoint list, database schema draft"],
        ["Sprint 1", "Backend Foundation", "Set up Flask app, SQLAlchemy ORM, user authentication (OTP via Fast2SMS), role-based access control, core API endpoints", "Working Flask backend with /register, /login, /verify-otp endpoints"],
        ["Sprint 2", "Disease Detection Module", "Train CNN (ResNet50) on PlantVillage dataset, implement Grad-CAM, build /detect-disease API endpoint, integrate with Flutter UI", "Disease detection screen with Grad-CAM heatmap overlay"],
        ["Sprint 3", "Soil Analysis Module", "Implement custom KNN algorithm (k=7), train on crop dataset, build /soil-analysis endpoint, create Flutter soil input form", "Soil analysis screen with crop recommendation and soil health score"],
        ["Sprint 4", "Weather & Marketplace", "Integrate OpenWeatherMap API, implement weather advisory logic, build crop listing/purchase/chat endpoints, Flutter marketplace screens", "Weather advisory screen, buyer/farmer marketplace with in-app chat"],
        ["Sprint 5", "Supporting Modules", "Implement AI chatbot (Groq LLaMA API), finance tracker, farming calendar, government schemes hub", "All 8 modules integrated and functional"],
        ["Sprint 6", "Integration & Testing", "System integration testing, API stress testing, UI/UX testing, bug fixes, performance optimisation", "Fully tested and integrated application"],
        ["Sprint 7", "Deployment & Validation", "Deploy backend on local server (host 0.0.0.0:5000), build Flutter APK, conduct end-to-end validation", "Production-ready APK and deployed Flask backend"],
    ],
    "Table 3.1 Mapping KrishiMitra Development to Agile Scrum Phases"
)

sec(doc, "3.3 V-Model Verification and Validation")
body(doc, "In addition to the Agile Scrum framework governing sprint planning and delivery, the KrishiMitra project applied V-Model principles for structured verification and validation at each development level. The V-Model (Verification and Validation model) is a software development lifecycle model where each development phase has a corresponding testing phase, forming a V-shaped process. The left arm of the V represents progressive decomposition from requirements to unit design, while the right arm represents the corresponding testing phases from unit testing to system acceptance testing [5]. This dual-framework approach — Agile for iterative delivery and V-Model for quality assurance — ensured that both speed of development and rigour of testing were maintained throughout the project.")
blank(doc)

body(doc, "The verification stages for KrishiMitra included: (1) Requirements Verification — confirming that the five project objectives were correctly captured as API specifications and Flutter screen requirements; (2) Architectural Verification — validating the three-tier client-server architecture against system design documents; (3) Module Design Verification — reviewing individual module designs (CNN pipeline, KNN algorithm, OTP flow) for logical correctness before implementation. The validation stages included: (1) Unit Testing — testing each Flask API endpoint independently using Postman with positive, negative, and boundary test cases; (2) Integration Testing — verifying that the Flutter frontend correctly consumed all 20+ backend API endpoints; (3) System Validation — end-to-end testing of complete user workflows (registration, disease detection, marketplace purchase) against the original user stories to confirm objective fulfilment [6].")
blank(doc)

para(doc, "Figure 3.2: V-Model applied to KrishiMitra — Verification and Validation phases", size=11, italic=True, align="CENTER")
blank(doc)

sec(doc, "3.4 Tools Used in Development")
body(doc, "The development of KrishiMitra utilised an integrated set of open-source and commercially available development tools. Visual Studio Code served as the primary Integrated Development Environment (IDE) for both Flutter (Dart) frontend development and Python Flask backend development, chosen for its lightweight footprint and extensive plugin ecosystem. Git version control was used for source code management, with commits structured per sprint to maintain a clear development history. The Flutter SDK (version 3.x) provided the cross-platform mobile development framework, enabling a single codebase to target both Android and iOS. Python 3.10 served as the backend runtime, with Flask, SQLAlchemy, TensorFlow, NumPy, and OpenCV as core dependencies. Postman was used for REST API testing and documentation. The PlantVillage dataset (available via TensorFlow Datasets) provided labelled crop disease images for CNN training [7].")
blank(doc)

br(doc)

# ═══════════════════════════════════════════════════
# CHAPTER 4: PROJECT MANAGEMENT
# ═══════════════════════════════════════════════════
chap(doc, "Chapter 4")
chap(doc, "Project Management")
blank(doc)

body(doc, "Effective project management is critical to the successful delivery of a complex, multi-module software system such as KrishiMitra. This chapter details the project timeline, risk analysis using PESTLE framework, and the estimated project budget. The project was managed using an Agile Scrum approach with clearly defined sprints, milestones, and deliverables. A Gantt chart was developed using Google Sheets to visually represent task scheduling, dependencies, and progress tracking across the seven sprints of the development lifecycle [8].")
blank(doc)

sec(doc, "4.1 Project Timeline")
sub(doc, "4.1.1 Project Planning Phase")
body(doc, "Table 4.1 summarises the timeline during the project planning phase. The planning phase spanned the first three weeks of the project and focused on requirement gathering, literature review, system architecture design, and technology stack selection. Key milestones during this phase included the finalisation of the five project objectives, the selection of Flutter and Flask as the core technology pair, and the design of the database schema. The planning phase concluded with a complete Product Backlog containing 47 user stories prioritised by business value and technical dependency.")
blank(doc)

add_table(doc,
    ["Task", "Week", "Start Date", "End Date", "Milestone"],
    [
        ["Requirement gathering & user story definition", "Week 1", "Jan 6, 2025", "Jan 10, 2025", "Product Backlog finalised"],
        ["Literature review (10 journal/conference papers)", "Week 1-2", "Jan 6, 2025", "Jan 17, 2025", "Literature Review complete"],
        ["System architecture design (3-tier)", "Week 2", "Jan 13, 2025", "Jan 17, 2025", "Architecture document approved"],
        ["Technology stack selection and setup", "Week 2", "Jan 13, 2025", "Jan 17, 2025", "Dev environment configured"],
        ["Database schema design (7 tables)", "Week 3", "Jan 20, 2025", "Jan 24, 2025", "ER diagram finalised"],
        ["API endpoint specification (20+ endpoints)", "Week 3", "Jan 20, 2025", "Jan 24, 2025", "API contract document ready"],
        ["Sprint planning (Sprints 1-7)", "Week 3", "Jan 22, 2025", "Jan 24, 2025", "Sprint backlog items assigned"],
    ],
    "Table 4.1 Project Planning Timeline"
)

sub(doc, "4.1.2 Project Implementation Phase")
body(doc, "Table 4.2 summarises the timeline during the project implementation phase. The implementation phase spanned fourteen weeks across seven sprints. Each sprint was two weeks in duration, with a one-week buffer built into the final sprint for integration testing and bug resolution. The Gantt chart clearly shows the sequential dependencies between modules — for example, the Flask backend authentication endpoints (Sprint 1) had to be completed before the marketplace module (Sprint 4) could begin, as the marketplace relies on authenticated user sessions. Similarly, the CNN model training (Sprint 2) was a prerequisite for the disease detection Flutter screen integration.")
blank(doc)

add_table(doc,
    ["Sprint", "Key Tasks", "Start", "End", "Deliverable"],
    [
        ["Sprint 1", "Flask setup, SQLAlchemy ORM, OTP auth, user registration/login endpoints", "Jan 27, 2025", "Feb 7, 2025", "Auth system live"],
        ["Sprint 2", "CNN training (ResNet50), Grad-CAM, /detect-disease endpoint, Flutter disease screen", "Feb 10, 2025", "Feb 21, 2025", "Disease detection module"],
        ["Sprint 3", "KNN algorithm, soil training data, /soil-analysis endpoint, Flutter soil screen", "Feb 24, 2025", "Mar 7, 2025", "Soil analysis module"],
        ["Sprint 4", "OpenWeatherMap integration, marketplace CRUD, in-app chat, Flutter screens", "Mar 10, 2025", "Mar 21, 2025", "Weather + Marketplace"],
        ["Sprint 5", "Groq LLaMA chatbot, finance tracker, farming calendar, govt schemes hub", "Mar 24, 2025", "Apr 4, 2025", "All 8 modules complete"],
        ["Sprint 6", "System integration, Postman API testing (42 test cases), UI/UX refinement", "Apr 7, 2025", "Apr 18, 2025", "Tested integrated system"],
        ["Sprint 7", "APK build, backend deployment, report writing, performance validation", "Apr 21, 2025", "May 2, 2025", "Final APK + report"],
    ],
    "Table 4.2 Project Implementation Timeline"
)

sec(doc, "4.2 Risk Analysis — PESTLE Framework")
body(doc, "PESTLE analysis is a strategic framework used to assess the Political, Economic, Social, Technological, Legal, and Environmental factors that may impact a project's success. Table 4.3 presents the PESTLE analysis for KrishiMitra, identifying key risks and corresponding mitigation strategies. Proactive risk identification allows the development team to design countermeasures into the system architecture and deployment plan, reducing the likelihood of project failure due to external environmental factors [9].")
blank(doc)

add_table(doc,
    ["Factor", "Risk", "Impact", "Likelihood", "Mitigation Strategy"],
    [
        ["Political", "Changes in government agricultural data sharing policies may restrict access to mandi price APIs or government scheme data", "High", "Low", "Use publicly available APIs (OpenWeatherMap, Agmarknet); design modular data sources that can be swapped"],
        ["Economic", "Farmers in target demographics may not afford smartphones with sufficient specs to run the Flutter app", "Medium", "Medium", "Optimise app for low-RAM devices (2GB+); test on Android Go devices; minimise APK size below 30MB"],
        ["Social", "Low digital literacy among rural farmers may hinder adoption of AI-driven features", "High", "High", "Design intuitive UI with icon-based navigation; implement Grad-CAM heatmaps for visual trust-building; plan for multi-language support"],
        ["Technological", "CNN model may underperform on unseen crop types or unusual disease presentations not in training data", "High", "Medium", "Use transfer learning (ResNet50 pretrained on ImageNet); apply data augmentation; clearly display confidence score to users"],
        ["Technological", "Groq API rate limits or downtime may disrupt chatbot functionality", "Medium", "Low", "Implement graceful error handling in Flutter; display fallback message; consider local LLM as future fallback"],
        ["Legal", "User phone numbers collected during OTP authentication are personal data subject to India's DPDPA 2023", "High", "High", "Store phone numbers hashed; implement OTP expiry (10 min); do not share data with third parties; add privacy policy screen"],
        ["Environmental", "Rural areas with poor 4G connectivity may experience timeouts on AI API calls", "Medium", "High", "Implement API timeout handling (10s); show loading indicators; cache last weather data locally; plan offline mode for future"],
    ],
    "Table 4.3 PESTLE Risk Analysis for KrishiMitra"
)

body(doc, "The most significant risks identified are in the Social and Technological categories. The social risk of low digital literacy is mitigated by KrishiMitra's design philosophy of icon-first navigation, colour-coded outputs (green for healthy, red for disease), and Grad-CAM visual explanations that do not require the farmer to understand the underlying AI model. The technological risk of connectivity is addressed through timeout handling and local caching of critical data such as the most recent weather advisory [10].")
blank(doc)

sec(doc, "4.3 Project Budget")
body(doc, "Table 4.4 presents the estimated project budget for KrishiMitra, covering human resources, software tools, cloud services, testing devices, and documentation. The budget reflects the open-source-first development philosophy adopted by the project, minimising licensing costs while maintaining professional-grade tooling. All AI/ML libraries (TensorFlow, scikit-learn equivalent functions, OpenCV) and development tools (Flutter SDK, Python, VS Code, Git) are free and open-source. The primary cost centres are the Fast2SMS OTP service (pay-per-use), the Groq API (free tier sufficient for development; paid tier for production scale), and the Android test device [11].")
blank(doc)

add_table(doc,
    ["Category", "Item", "Unit Cost (INR)", "Quantity", "Total Cost (INR)", "Notes"],
    [
        ["Human Resources", "Developer (self, student project)", "0", "1", "0", "Student project — no paid labour"],
        ["Software Tools", "Flutter SDK, Python, VS Code, Git", "0", "—", "0", "All open-source and free"],
        ["AI/ML Libraries", "TensorFlow, NumPy, OpenCV, Flask", "0", "—", "0", "Open-source Python packages"],
        ["API Services", "Fast2SMS OTP (pay-per-use)", "0.20/SMS", "500 SMS", "100", "Testing and demo phase"],
        ["API Services", "OpenWeatherMap API (free tier)", "0", "—", "0", "Free up to 1000 calls/day"],
        ["API Services", "Groq API (free tier)", "0", "—", "0", "Free tier: 14,400 requests/day"],
        ["Dataset", "PlantVillage (TensorFlow Datasets)", "0", "—", "0", "Open-access academic dataset"],
        ["Hardware", "Android test device (mid-range)", "12,000", "1", "12,000", "For APK testing on physical device"],
        ["Documentation", "Report printing and binding", "500", "2 copies", "1,000", "Final project report submission"],
        ["Contingency", "Miscellaneous (10%)", "—", "—", "1,310", "10% of total non-zero costs"],
        ["", "Total Estimated Budget", "", "", "14,410 INR", "≈ USD 173"],
    ],
    "Table 4.4 Estimated Project Budget for KrishiMitra"
)

body(doc, "The total estimated project budget for KrishiMitra is approximately INR 14,410 (USD 173), making it an exceptionally cost-effective AI-powered agricultural platform. The negligible software cost is a direct result of the open-source technology stack selected during the planning phase. The primary hardware investment is a mid-range Android device for physical testing, ensuring that the application performs correctly on devices representative of the target farmer demographic. The Fast2SMS OTP service cost of INR 100 covers approximately 500 OTP messages, sufficient for the development, testing, and demonstration phases of the project [12].")
blank(doc)

# References
br(doc)
sec(doc, "References for Chapters 3 and 4")
refs = [
    "[1] Beck, K., et al. (2001). Manifesto for Agile Software Development. Agile Alliance. Available: https://agilemanifesto.org",
    "[2] Schwaber, K. and Sutherland, J. (2020). The Scrum Guide. Scrum.org. Available: https://scrumguides.org",
    "[3] Rubin, K. S. (2012). Essential Scrum: A Practical Guide to the Most Popular Agile Process. Addison-Wesley Professional.",
    "[4] Cohn, M. (2004). User Stories Applied: For Agile Software Development. Addison-Wesley Professional.",
    "[5] Forsberg, K. and Mooz, H. (1991). The relationship of system engineering to the project cycle. Proceedings of the First Annual Symposium of National Council on Systems Engineering, pp. 57–65.",
    "[6] Pressman, R. S. and Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach, 9th ed. McGraw-Hill Education.",
    "[7] Hughes, D. P. and Salathe, M. (2015). An open access repository of images on plant health to enable the development of mobile disease diagnostics. arXiv preprint arXiv:1511.08060.",
    "[8] Project Management Institute (2021). A Guide to the Project Management Body of Knowledge (PMBOK Guide), 7th ed. PMI.",
    "[9] Yüksel, I. (2012). Developing a multi-criteria decision making model for PESTEL analysis. International Journal of Business and Management, 7(24), pp. 52–66.",
    "[10] Mohanty, S. P., Hughes, D. P. and Salathe, M. (2016). Using deep learning for image-based plant disease detection. Frontiers in Plant Science, 7, p. 1419.",
    "[11] Aggarwal, P. K., et al. (2022). Bringing digital revolution to Indian smallholder farmers. Nature Food, 3, pp. 14–24.",
    "[12] NITI Aayog (2017). Doubling Farmers' Income: Rationale, Strategy, Prospects and Action Plan. Government of India, New Delhi.",
]
for r in refs:
    body(doc, r)

doc.save("KrishiMitra_Ch3_Ch4.docx")
print("Chapters 3 & 4 saved to KrishiMitra_Ch3_Ch4.docx")
