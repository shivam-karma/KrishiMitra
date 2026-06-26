from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_height=Cm(29.7); s.page_width=Cm(21.0)
        s.left_margin=Cm(2.54); s.right_margin=Cm(2.54)
        s.top_margin=Cm(2.54); s.bottom_margin=Cm(2.54)
    return doc

def para(doc, text, size=12, bold=False, align="JUSTIFY", italic=False):
    p = doc.add_paragraph()
    p.alignment = {"CENTER":WD_ALIGN_PARAGRAPH.CENTER,"LEFT":WD_ALIGN_PARAGRAPH.LEFT}.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.line_spacing = Pt(21)
    r = p.add_run(text); r.bold=bold; r.italic=italic
    r.font.size=Pt(size); r.font.name="Times New Roman"
    return p

def chap(d,t): para(d,t,16,True,"CENTER")
def sec(d,t):  para(d,t,14,True,"LEFT")
def sub(d,t):  para(d,t,12,True,"LEFT")
def body(d,t): para(d,t,12,align="JUSTIFY")
def cap(d,t):  para(d,t,11,False,"CENTER",True)
def blank(d):  para(d,"",align="LEFT")
def br(d):     d.add_page_break()

def tbl(doc, headers, rows, caption=""):
    if caption:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(caption); r.bold=True; r.font.size=Pt(11); r.font.name="Times New Roman"
    t=doc.add_table(rows=1+len(rows), cols=len(headers)); t.style="Table Grid"
    hr=t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h
        for run in c.paragraphs[0].runs:
            run.bold=True; run.font.size=Pt(10); run.font.name="Times New Roman"
            run.font.color.rgb=RGBColor(255,255,255)
        tc=c._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'1F3864'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]; fill='EEF4FF' if ri%2==0 else 'FFFFFF'
        for ci,val in enumerate(row):
            c=tr.cells[ci]; c.text=str(val)
            for p in c.paragraphs:
                for run in p.runs: run.font.size=Pt(10); run.font.name="Times New Roman"
            tc=c._tc; tcPr=tc.get_or_add_tcPr()
            shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); shd.set(qn('w:val'),'clear'); tcPr.append(shd)
    blank(doc)

doc = new_doc()
chap(doc,"Chapter 6")
chap(doc,"Hardware, Software and Simulation")
blank(doc)

body(doc,"Chapter 6 details the complete implementation of the KrishiMitra platform — covering the software development tools used, the core algorithmic code for each module with commentary, and the simulation and testing environment. Since KrishiMitra is a software-only system (mobile application + REST API backend), the 'hardware' component refers to the physical devices used during development and testing: a developer laptop and an Android smartphone for APK testing. There is no custom electronic hardware or embedded microcontroller involved. The implementation follows the sprint plan defined in Chapter 4, with each module developed, unit-tested, and integrated incrementally [1].")
blank(doc)

sec(doc,"6.1 Development Hardware Configuration")
body(doc,"Table 6.1 lists the hardware configuration used during the development and testing of KrishiMitra. The development laptop hosted the Flask backend, SQLite database, and model training pipeline. The Android test device was used to install and validate the Flutter APK across real-world network conditions, camera functionality, and touch interactions.")
blank(doc)

tbl(doc,
    ["Component","Specification","Purpose"],
    [
        ["Development Laptop","Intel Core i5-11th Gen, 8GB RAM, 512GB SSD, Windows 11","Flask backend development, CNN model training, Flutter build"],
        ["Android Test Device","Realme C55, Android 13, 6GB RAM, 128GB storage, 50MP camera","Flutter APK testing, camera-based disease detection validation"],
        ["Network Environment","Home Wi-Fi (50 Mbps) + mobile 4G hotspot (15 Mbps)","Testing API response times under different connectivity conditions"],
        ["Local Server","Flask dev server on localhost:5000 (developer laptop)","Backend API hosting during development and testing phases"],
    ],
    "Table 6.1 Development and Testing Hardware Configuration"
)

sec(doc,"6.2 Software Development Tools")
body(doc,"Table 6.2 describes the software tools used across the KrishiMitra development lifecycle. Each tool was selected for its suitability to the specific development task, open-source availability, and compatibility with the project's technology stack.")
blank(doc)

tbl(doc,
    ["Tool","Version","Category","Purpose in KrishiMitra"],
    [
        ["Visual Studio Code","1.88+","IDE","Primary code editor for Dart/Flutter and Python/Flask development; extensions: Flutter, Dart, Python, REST Client"],
        ["Flutter SDK","3.19+","Mobile Framework","Cross-platform mobile app development; hot reload for rapid UI iteration"],
        ["Android Studio","Hedgehog+","Android Build","Android SDK management, AVD emulator for device simulation, APK build signing"],
        ["Python","3.11","Backend Runtime","Flask server, TensorFlow model training, KNN algorithm, data preprocessing"],
        ["Flask","3.0+","Web Framework","REST API server with 20+ endpoints; Flask-CORS for cross-origin support"],
        ["SQLAlchemy","2.0+","ORM","Database schema definition, migration, and query abstraction over SQLite"],
        ["TensorFlow / Keras","2.15+","ML Framework","CNN (ResNet50) model training, Grad-CAM implementation, model serialisation"],
        ["OpenCV","4.9+","Image Processing","Image preprocessing (resize, normalise) before CNN inference; Grad-CAM overlay"],
        ["NumPy","1.26+","Numerical Computing","Array operations for KNN distance calculations and image pixel manipulation"],
        ["Postman","11+","API Testing","REST API endpoint testing with automated test collections; 42 test cases"],
        ["Git / GitHub","2.44+","Version Control","Source code management, sprint-based commit history, code backup"],
        ["Google Sheets","Web","Gantt Chart","Project timeline and sprint planning Gantt chart visualisation"],
        ["Draw.io (diagrams.net)","Web","Diagram Tool","System architecture, DFD, use case, flowchart diagram creation"],
    ],
    "Table 6.2 Software Development Tools Used in KrishiMitra"
)

sec(doc,"6.3 Software Code — Core Modules")
sub(doc,"6.3.1 Flask Application Entry Point")
body(doc,"The Flask application is initialised in app.py. The following code snippet shows the application factory pattern used to create the Flask app instance, configure CORS, initialise the SQLAlchemy ORM, and register all blueprint routes.")
blank(doc)

code1 = [
    "# app.py — KrishiMitra Flask Backend Entry Point",
    "from flask import Flask",
    "from flask_sqlalchemy import SQLAlchemy",
    "from flask_cors import CORS",
    "",
    "db = SQLAlchemy()  # SQLAlchemy instance — shared across all models",
    "",
    "def create_app():",
    "    app = Flask(__name__)",
    "    # Configure SQLite database path",
    "    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///krishimitra.db'",
    "    app.config['SECRET_KEY'] = 'krishimitra-secret-2025'",
    "    CORS(app)          # Enable CORS for Flutter client requests",
    "    db.init_app(app)   # Bind SQLAlchemy to this app instance",
    "",
    "    with app.app_context():",
    "        db.create_all()  # Create all tables if they do not exist",
    "",
    "    # Register API blueprints (routes)",
    "    from routes.auth import auth_bp",
    "    from routes.disease import disease_bp",
    "    from routes.soil import soil_bp",
    "    from routes.marketplace import market_bp",
    "    app.register_blueprint(auth_bp, url_prefix='/auth')",
    "    app.register_blueprint(disease_bp, url_prefix='/disease')",
    "    app.register_blueprint(soil_bp, url_prefix='/soil')",
    "    app.register_blueprint(market_bp, url_prefix='/marketplace')",
    "    return app",
    "",
    "if __name__ == '__main__':",
    "    app = create_app()",
    "    app.run(host='0.0.0.0', port=5000, debug=True)  # Listen on all interfaces",
]
p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(16)
r = p.add_run("\n".join(code1))
r.font.name = "Courier New"; r.font.size = Pt(9)

blank(doc)
body(doc,"The application uses Flask's blueprint pattern to organise routes into logical modules (auth, disease, soil, marketplace). The CORS extension is applied globally to allow the Flutter mobile client running on a different origin (Android device) to communicate with the Flask server. The SQLite database file (krishimitra.db) is automatically created in the instance folder when the application starts for the first time [2].")
blank(doc)

sub(doc,"6.3.2 CNN Disease Detection Module")
body(doc,"The disease detection module loads the pre-trained ResNet50 model and applies Grad-CAM to produce a visual heatmap alongside the classification result. The following code shows the inference pipeline.")
blank(doc)

code2 = [
    "# routes/disease.py — CNN Disease Detection with Grad-CAM",
    "import tensorflow as tf",
    "import numpy as np, cv2, base64",
    "from flask import Blueprint, request, jsonify",
    "",
    "disease_bp = Blueprint('disease', __name__)",
    "MODEL_PATH = 'models/disease_model.h5'",
    "CLASSES = ['Healthy', 'Leaf_Spot', 'Rust']  # 3 disease classes",
    "model = tf.keras.models.load_model(MODEL_PATH)  # Load saved CNN model",
    "",
    "def preprocess(img_bytes):",
    "    # Decode image bytes, resize to 224x224 (ResNet50 input size), normalise",
    "    nparr = np.frombuffer(img_bytes, np.uint8)",
    "    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)",
    "    img = cv2.resize(img, (224, 224))",
    "    img = img.astype('float32') / 255.0  # Normalise pixel values to [0,1]",
    "    return np.expand_dims(img, axis=0)   # Add batch dimension",
    "",
    "def grad_cam(model, img_array, class_idx):",
    "    # Get last convolutional layer for gradient computation",
    "    grad_model = tf.keras.Model(inputs=model.inputs,",
    "        outputs=[model.get_layer('conv5_block3_out').output, model.output])",
    "    with tf.GradientTape() as tape:",
    "        conv_outputs, predictions = grad_model(img_array)",
    "        loss = predictions[:, class_idx]  # Focus on predicted class",
    "    grads = tape.gradient(loss, conv_outputs)  # Compute gradients",
    "    pooled = tf.reduce_mean(grads, axis=(0,1,2))  # Global average pooling",
    "    cam = conv_outputs[0] @ pooled[..., tf.newaxis]",
    "    cam = tf.squeeze(cam).numpy()",
    "    cam = np.maximum(cam, 0)  # ReLU — only positive activations",
    "    cam = cam / cam.max()     # Normalise to [0,1]",
    "    return cv2.resize(cam, (224, 224))  # Resize heatmap to input size",
    "",
    "@disease_bp.route('/detect', methods=['POST'])",
    "def detect_disease():",
    "    if 'image' not in request.files:",
    "        return jsonify({'status': 'error', 'message': 'No image provided'}), 400",
    "    img_bytes = request.files['image'].read()",
    "    img_array = preprocess(img_bytes)  # Preprocess image",
    "    preds = model.predict(img_array)   # Run CNN inference",
    "    class_idx = int(np.argmax(preds))  # Get predicted class index",
    "    confidence = float(np.max(preds))  # Get confidence score",
    "    heatmap = grad_cam(model, img_array, class_idx)  # Compute Grad-CAM",
    "    # Overlay heatmap on original image",
    "    heatmap_colored = cv2.applyColorMap((heatmap*255).astype(np.uint8), cv2.COLORMAP_JET)",
    "    orig = cv2.imdecode(np.frombuffer(img_bytes, np.uint8), cv2.IMREAD_COLOR)",
    "    orig = cv2.resize(orig, (224, 224))",
    "    overlay = cv2.addWeighted(orig, 0.6, heatmap_colored, 0.4, 0)",
    "    _, buffer = cv2.imencode('.jpg', overlay)",
    "    heatmap_b64 = base64.b64encode(buffer).decode()  # Encode as base64 string",
    "    return jsonify({",
    "        'status': 'success',",
    "        'disease': CLASSES[class_idx],    # Predicted disease class",
    "        'confidence': round(confidence*100, 2),  # Confidence as percentage",
    "        'heatmap': heatmap_b64            # Grad-CAM overlay image",
    "    })",
]
p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(16)
r = p.add_run("\n".join(code2))
r.font.name = "Courier New"; r.font.size = Pt(8)

blank(doc)
body(doc,"The detect_disease endpoint accepts a multipart form-data POST request with the crop leaf image. The image is preprocessed to 224x224 pixels and normalised. The ResNet50 model performs forward-pass inference and returns class probabilities. The Grad-CAM function computes the gradient of the top class score with respect to the last convolutional layer's feature maps, producing a heatmap that highlights the regions most responsible for the prediction. The heatmap is overlaid on the original image and returned as a base64-encoded JPEG string, which the Flutter client decodes and displays directly in the Image widget [3].")
blank(doc)

sub(doc,"6.3.3 KNN Soil Analysis Module")
body(doc,"The soil analysis module implements a K-Nearest Neighbours algorithm from scratch in pure Python, without any dependency on scikit-learn. This design choice ensures maximum deployment flexibility across different server environments.")
blank(doc)

code3 = [
    "# routes/soil.py — Custom KNN Crop Recommendation Engine",
    "import json, math",
    "from flask import Blueprint, request, jsonify",
    "",
    "soil_bp = Blueprint('soil', __name__)",
    "",
    "# Load training data: list of {N, P, K, pH, moisture, crop} dicts",
    "with open('data/soil_data.json') as f:",
    "    TRAINING_DATA = json.load(f)",
    "",
    "def euclidean_distance(a, b):",
    "    # Compute Euclidean distance between two 5-feature vectors",
    "    return math.sqrt(sum((a[i]-b[i])**2 for i in range(len(a))))",
    "",
    "def knn_predict(features, k=7):",
    "    # Compute distance from input to every training sample",
    "    distances = []",
    "    for sample in TRAINING_DATA:",
    "        train_feat = [sample['N'], sample['P'], sample['K'],",
    "                      sample['pH'], sample['moisture']]",
    "        dist = euclidean_distance(features, train_feat)",
    "        distances.append((dist, sample['crop']))",
    "    distances.sort(key=lambda x: x[0])  # Sort by ascending distance",
    "    k_nearest = distances[:k]            # Select k nearest neighbours",
    "    # Majority vote: count occurrences of each crop label",
    "    votes = {}",
    "    for dist, crop in k_nearest:",
    "        votes[crop] = votes.get(crop, 0) + 1",
    "    predicted_crop = max(votes, key=votes.get)",  
    "    confidence = (votes[predicted_crop] / k) * 100  # Confidence as %",
    "    return predicted_crop, confidence",
    "",
    "def soil_health_score(N, P, K, pH, moisture):",
    "    # Compute a composite soil health score (0-100)",
    "    n_score = min(N/140, 1) * 25   # N optimal up to 140 kg/ha",
    "    p_score = min(P/60, 1) * 25    # P optimal up to 60 kg/ha",
    "    k_score = min(K/200, 1) * 25   # K optimal up to 200 kg/ha",
    "    ph_score = max(0, 1 - abs(pH-6.5)/3.5) * 25  # Ideal pH = 6.5",
    "    return round(n_score + p_score + k_score + ph_score, 1)",
    "",
    "@soil_bp.route('/analyse', methods=['POST'])",
    "def analyse_soil():",
    "    data = request.get_json()",
    "    try:",
    "        # Extract soil parameter inputs from request body",
    "        N = float(data['nitrogen'])       # Nitrogen in kg/ha",
    "        P = float(data['phosphorus'])     # Phosphorus in kg/ha",
    "        K = float(data['potassium'])      # Potassium in kg/ha",
    "        pH = float(data['ph'])            # Soil pH (3.0 - 10.0)",
    "        moisture = float(data['moisture'])# Moisture percentage (0-100)",
    "    except (KeyError, TypeError):",
    "        return jsonify({'status':'error','message':'Invalid input'}), 400",
    "    features = [N, P, K, pH, moisture]",
    "    crop, confidence = knn_predict(features, k=7)  # Run KNN prediction",
    "    health_score = soil_health_score(N, P, K, pH, moisture)",
    "    # Generate fertilizer recommendation based on deficiencies",
    "    recommendations = []",
    "    if N < 40: recommendations.append('Apply Urea: 50 kg/ha to boost Nitrogen')",
    "    if P < 20: recommendations.append('Apply DAP: 30 kg/ha to boost Phosphorus')",
    "    if K < 80: recommendations.append('Apply MOP: 40 kg/ha to boost Potassium')",
    "    if pH < 5.5: recommendations.append('Apply lime to raise soil pH')",
    "    if pH > 8.0: recommendations.append('Apply gypsum to lower soil pH')",
    "    return jsonify({",
    "        'status': 'success',",
    "        'recommended_crop': crop,         # Best crop for these soil params",
    "        'confidence': round(confidence,1),# Prediction confidence %",
    "        'soil_health_score': health_score,# Composite health score (0-100)",
    "        'fertilizer_advice': recommendations  # List of actionable advice",
    "    })",
]
p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(16)
r = p.add_run("\n".join(code3))
r.font.name = "Courier New"; r.font.size = Pt(8)

blank(doc)
body(doc,"The KNN implementation requires no external machine learning libraries. The euclidean_distance function computes the L2 distance between the input feature vector [N, P, K, pH, moisture] and each training sample. The k=7 nearest samples are selected and a majority vote determines the recommended crop. The soil_health_score function computes a composite 0-100 score by normalising each soil parameter against its agronomic optimum and summing the weighted subscores. Fertiliser recommendations are generated using rule-based threshold logic applied to the raw input values [4].")
blank(doc)

sub(doc,"6.3.4 Flutter Disease Detection Screen")
body(doc,"The following Dart code illustrates the core logic of the DiseaseDetectionScreen in the Flutter frontend, showing how the image is captured, sent to the Flask API, and the Grad-CAM result is rendered.")
blank(doc)

code4 = [
    "// disease_detection_screen.dart — Flutter Disease Detection UI",
    "import 'dart:convert'; import 'dart:io';",
    "import 'package:flutter/material.dart';",
    "import 'package:image_picker/image_picker.dart';",
    "import 'package:http/http.dart' as http;",
    "",
    "class DiseaseDetectionScreen extends StatefulWidget {",
    "  @override State<DiseaseDetectionScreen> createState()",
    "    => _DiseaseDetectionScreenState();",
    "}",
    "",
    "class _DiseaseDetectionScreenState extends State<DiseaseDetectionScreen> {",
    "  File? _image;         // Selected/captured image file",
    "  String? _disease;     // Predicted disease class from API",
    "  double? _confidence;  // Prediction confidence percentage",
    "  String? _heatmapB64;  // Grad-CAM heatmap as base64 string",
    "  bool _isLoading = false;",
    "",
    "  // Open camera to capture a leaf photograph",
    "  Future<void> _captureImage() async {",
    "    final picker = ImagePicker();",
    "    final picked = await picker.pickImage(source: ImageSource.camera);",
    "    if (picked != null) setState(() => _image = File(picked.path));",
    "  }",
    "",
    "  // Submit image to Flask /disease/detect endpoint",
    "  Future<void> _analyseDisease() async {",
    "    if (_image == null) return;",
    "    setState(() => _isLoading = true);  // Show loading indicator",
    "    var request = http.MultipartRequest('POST',",
    "      Uri.parse('http://192.168.1.100:5000/disease/detect'));",
    "    request.files.add(await http.MultipartFile.fromPath('image', _image!.path));",
    "    final response = await request.send();",
    "    final body = await response.stream.bytesToString();",
    "    final data = jsonDecode(body);",
    "    setState(() {",
    "      _disease = data['disease'];         // e.g. 'Leaf_Spot'",
    "      _confidence = data['confidence'];   // e.g. 87.3",
    "      _heatmapB64 = data['heatmap'];      // base64 Grad-CAM image",
    "      _isLoading = false;",
    "    });",
    "  }",
    "",
    "  @override Widget build(BuildContext context) {",
    "    return Scaffold(",
    "      appBar: AppBar(title: Text('Disease Detection')),",
    "      body: Column(children: [",
    "        // Display selected image or placeholder",
    "        _image != null ? Image.file(_image!, height: 200) : Icon(Icons.camera_alt, size: 80),",
    "        ElevatedButton(onPressed: _captureImage, child: Text('Capture Leaf Image')),",
    "        ElevatedButton(onPressed: _analyseDisease, child: Text('Analyse Disease')),",
    "        if (_isLoading) CircularProgressIndicator(),",
    "        // Display Grad-CAM heatmap if available",
    "        if (_heatmapB64 != null)",
    "          Image.memory(base64Decode(_heatmapB64!), height: 200),",
    "        if (_disease != null)",
    "          Text('Diagnosis: $_disease (${_confidence?.toStringAsFixed(1)}%)',",
    "            style: TextStyle(fontSize: 18, fontWeight: FontWeight.bold)),",
    "      ]),",
    "    );",
    "  }",
    "}",
]
p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(16)
r = p.add_run("\n".join(code4))
r.font.name = "Courier New"; r.font.size = Pt(8)

blank(doc)
body(doc,"The DiseaseDetectionScreen uses the image_picker package to access the device camera and capture a leaf image. The image is sent as a multipart HTTP POST request to the Flask /disease/detect endpoint. The API response is decoded from JSON, and the Grad-CAM heatmap (returned as a base64 string) is decoded and displayed using Image.memory(). A CircularProgressIndicator is shown during the API call to provide feedback to the user during the inference period [5].")
blank(doc)

sec(doc,"6.4 Simulation and Testing Environment")
body(doc,"Since KrishiMitra is a software system, simulation takes the form of virtual device emulation and API mocking rather than physical circuit simulation. The following simulation and testing tools were used:")
blank(doc)

tbl(doc,
    ["Tool","Type","Purpose"],
    [
        ["Android Emulator (AVD Manager)","Mobile Simulator","Simulates Android devices of various screen sizes and Android versions without a physical device; used for UI layout testing"],
        ["Postman","API Simulator","Sends HTTP requests to Flask endpoints with controlled inputs; verifies response schemas, status codes, and error handling; 42 test cases executed"],
        ["TensorFlow Model Evaluate","ML Simulator","Evaluates CNN model on validation dataset split to measure accuracy, precision, recall, and F1 score before deployment"],
        ["Flask Test Client","Backend Simulator","Python unittest-based Flask test client that sends in-process HTTP requests to API endpoints without a network connection"],
        ["Chrome DevTools (Network tab)","Network Simulator","Simulates slow 3G/4G connections to test Flutter app behaviour under constrained bandwidth conditions"],
        ["Wokwi (reference only)","Circuit Simulator","Not directly used (no embedded hardware); referenced for understanding IoT integration concepts for future soil sensor work"],
    ],
    "Table 6.3 Simulation and Testing Tools"
)

body(doc,"The primary simulation environment for KrishiMitra is the combination of the Android Virtual Device (AVD) emulator for UI testing and Postman for API contract validation. The CNN model was evaluated on a held-out test split (20% of the PlantVillage dataset) using TensorFlow's model.evaluate() function, producing the accuracy and loss metrics reported in Chapter 7. The Flask test client was used for automated regression testing of all API endpoints, ensuring that code changes in later sprints did not break previously validated functionality [6].")
blank(doc)

br(doc)
sec(doc,"References for Chapter 6")
for r in [
    "[1] Lutz, M. (2013). Learning Python, 5th ed. O'Reilly Media.",
    "[2] Grinberg, M. (2018). Flask Web Development, 2nd ed. O'Reilly Media.",
    "[3] Selvaraju, R. R., et al. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. Proceedings of the IEEE International Conference on Computer Vision (ICCV), pp. 618–626.",
    "[4] Pudumalar, S., et al. (2017). Crop recommendation system for precision agriculture. Proceedings of the 8th IEEE International Conference on Computing, Communication and Networking Technologies (ICCCNT).",
    "[5] Google LLC (2024). Flutter HTTP package documentation. pub.dev. Available: https://pub.dev/packages/http",
    "[6] Microsoft (2023). Visual Studio Code documentation. Available: https://code.visualstudio.com/docs",
]:
    body(doc,r)

doc.save("KrishiMitra_Ch6.docx")
print("Chapter 6 saved.")
