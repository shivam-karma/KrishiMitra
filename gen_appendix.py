from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_height=Cm(29.7); s.page_width=Cm(21.0)
        s.left_margin=Cm(2.54); s.right_margin=Cm(2.54)
        s.top_margin=Cm(2.54); s.bottom_margin=Cm(2.54)
    return doc

def title(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(18)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(16); r.font.name="Times New Roman"

def sec(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.name="Times New Roman"

def sub(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.name="Times New Roman"

def body(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing=Pt(21); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.font.size=Pt(12); r.font.name="Times New Roman"

def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing=Pt(18)
    r=p.add_run(text); r.font.size=Pt(11); r.font.name="Times New Roman"

def cap(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.name="Times New Roman"

def blank(doc): para=doc.add_paragraph(); para.paragraph_format.line_spacing=Pt(12)
def br(doc): doc.add_page_break()

def tbl(doc, headers, rows, caption="", widths=None):
    if caption:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(caption); r.bold=True; r.font.size=Pt(11); r.font.name="Times New Roman"
    t=doc.add_table(rows=1+len(rows), cols=len(headers)); t.style="Table Grid"
    hr=t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h
        for run in c.paragraphs[0].runs:
            run.bold=True; run.font.size=Pt(10); run.font.name="Times New Roman"
            run.font.color.rgb=RGBColor(255,255,255)
        tc=c._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'1F3864'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]; fill='EEF4FF' if ri%2==0 else 'FFFFFF'
        for ci,val in enumerate(row):
            c=tr.cells[ci]; c.text=str(val)
            for p in c.paragraphs:
                for run in p.runs: run.font.size=Pt(10); run.font.name="Times New Roman"
            tc=c._tc; tcPr=tc.get_or_add_tcPr()
            shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); shd.set(qn('w:val'),'clear'); tcPr.append(shd)
    if widths:
        for row in t.rows:
            for i,cell in enumerate(row.cells):
                if i < len(widths): cell.width=widths[i]
    blank(doc)

# ════════════════════════════════════════════
doc = new_doc()
title(doc, "Appendix")
blank(doc)

body(doc, "This appendix provides supplementary material for the KrishiMitra Major Project Report, including specifications of the software libraries and datasets used (Appendix A), the complete API endpoint reference (Appendix B), a summary of the training datasets (Appendix C), project screenshot descriptions demonstrating key scenarios (Appendix D), and installation and configuration instructions (Appendix E).")
blank(doc)

# ════════════════════════════════════════════
# APPENDIX A — Software Library Specifications (Datasheets)
# ════════════════════════════════════════════
br(doc)
sec(doc, "Appendix A: Software Library Specifications")
body(doc, "The following tables summarise the specifications of the key software libraries and frameworks used in the KrishiMitra project. These serve as the equivalent of component datasheets for a software-based system, documenting version, purpose, key functions, and configuration details.")
blank(doc)

sub(doc, "A.1 TensorFlow / Keras — CNN Model Library")
tbl(doc,
    ["Specification", "Details"],
    [
        ["Library Name", "TensorFlow with Keras high-level API"],
        ["Version Used", "TensorFlow 2.15.0 / Keras 2.15.0"],
        ["Developer", "Google Brain Team / TensorFlow Contributors"],
        ["Licence", "Apache License 2.0 (Open Source, Free)"],
        ["Purpose in KrishiMitra", "Training and inference of the ResNet50 CNN model for crop disease detection; Grad-CAM heatmap computation"],
        ["Key Functions Used", "tf.keras.applications.ResNet50(), model.predict(), model.fit(), GradientTape(), tf.keras.models.load_model()"],
        ["Model Architecture", "ResNet50 — 50-layer deep residual network with skip connections; pretrained on ImageNet (1000 classes)"],
        ["Input Specification", "4D tensor: (batch_size, 224, 224, 3) — RGB image normalised to [0,1]"],
        ["Output Specification", "1D probability vector: shape (batch_size, num_classes) — softmax probabilities per disease class"],
        ["Training Configuration", "Optimiser: Adam (lr=0.001); Loss: Categorical Cross-Entropy; Epochs: 25; Batch Size: 32"],
        ["Model File Format", "TensorFlow SavedModel (.h5 format) — portable and version-controlled"],
        ["Hardware Requirement", "CPU-only training feasible (slow); GPU (NVIDIA CUDA) recommended for training >10 epochs"],
        ["Installation", "pip install tensorflow==2.15.0"],
        ["Documentation URL", "https://www.tensorflow.org/api_docs"],
    ],
    "Table A.1 TensorFlow/Keras Library Specification",
    [Cm(6), Cm(11)]
)

sub(doc, "A.2 Flask — REST API Web Framework")
tbl(doc,
    ["Specification", "Details"],
    [
        ["Library Name", "Flask"],
        ["Version Used", "Flask 3.0.2"],
        ["Developer", "Pallets Projects (Armin Ronacher)"],
        ["Licence", "BSD 3-Clause (Open Source, Free)"],
        ["Purpose in KrishiMitra", "REST API server hosting all 20+ backend endpoints; blueprint-based route organisation; request/response handling"],
        ["Key Extensions Used", "Flask-CORS (cross-origin requests from Flutter), Flask-SQLAlchemy (ORM integration)"],
        ["Key Decorators Used", "@app.route(), @blueprint.route(), request.get_json(), request.files, jsonify()"],
        ["Server Configuration", "host='0.0.0.0', port=5000, debug=True (development); Gunicorn for production"],
        ["Request Handling", "Synchronous (WSGI); each request processed sequentially per worker thread"],
        ["Error Handling", "HTTP 400 (bad input), 401 (unauthorised), 404 (not found), 500 (server error) with JSON error body"],
        ["CORS Configuration", "CORS(app) — allows all origins; restrict to Flutter app IP in production"],
        ["Installation", "pip install flask flask-cors flask-sqlalchemy"],
        ["Documentation URL", "https://flask.palletsprojects.com"],
    ],
    "Table A.2 Flask Framework Specification",
    [Cm(6), Cm(11)]
)

sub(doc, "A.3 Flutter SDK — Mobile Frontend Framework")
tbl(doc,
    ["Specification", "Details"],
    [
        ["Framework Name", "Flutter"],
        ["Version Used", "Flutter 3.19.x / Dart 3.3.x"],
        ["Developer", "Google LLC"],
        ["Licence", "BSD 3-Clause (Open Source, Free)"],
        ["Purpose in KrishiMitra", "Cross-platform mobile UI — single codebase for Android (and iOS); all 15 application screens"],
        ["Rendering Engine", "Skia / Impeller — custom rendering (not native UI widgets); ensures pixel-perfect cross-platform UI"],
        ["Target Platform", "Android APK (primary); iOS IPA (future)"],
        ["Min Android SDK", "API Level 21 (Android 5.0 Lollipop)"],
        ["Key Packages Used", "http: ^1.2.0 (API calls), image_picker: ^1.0.7 (camera), shared_preferences: ^2.2.2 (session), provider: ^6.1.2 (state management)"],
        ["Navigation", "GoRouter v14 — declarative routing with path-based navigation and role guards"],
        ["State Management", "Provider pattern — ChangeNotifier for reactive UI updates"],
        ["Build Command (APK)", "flutter build apk --release"],
        ["Installation", "Download Flutter SDK from https://flutter.dev/docs/get-started/install"],
        ["Documentation URL", "https://flutter.dev/docs"],
    ],
    "Table A.3 Flutter SDK Specification",
    [Cm(6), Cm(11)]
)

sub(doc, "A.4 SQLite / SQLAlchemy — Database")
tbl(doc,
    ["Specification", "Details"],
    [
        ["Database Engine", "SQLite 3.x"],
        ["ORM Library", "SQLAlchemy 2.0.x"],
        ["Licence", "SQLite: Public Domain; SQLAlchemy: MIT (Open Source, Free)"],
        ["Purpose in KrishiMitra", "Persistent data storage for users, listings, purchases, messages, finance records, and calendar tasks"],
        ["Database File", "krishimitra.db — single file, auto-created on first run"],
        ["Tables", "7 tables: users, listings, purchases, messages, finance_records, calendar_tasks, soil_results"],
        ["Max Concurrent Connections", "SQLite supports multiple readers; single writer (WAL mode recommended for production)"],
        ["Data Types Used", "Integer, String, Float, Boolean, DateTime, Text, Date"],
        ["Schema Management", "db.create_all() — auto-creates all tables from SQLAlchemy model definitions"],
        ["Backup", "Single file copy — krishimitra.db is fully portable"],
        ["Installation", "pip install sqlalchemy (SQLite built into Python standard library)"],
        ["Documentation URL", "https://docs.sqlalchemy.org / https://www.sqlite.org"],
    ],
    "Table A.4 SQLite/SQLAlchemy Database Specification",
    [Cm(6), Cm(11)]
)

br(doc)

# ════════════════════════════════════════════
# APPENDIX B — Complete API Endpoint Reference
# ════════════════════════════════════════════
sec(doc, "Appendix B: Complete API Endpoint Reference")
body(doc, "Table B.1 lists all REST API endpoints implemented in the KrishiMitra Flask backend. Each endpoint is specified with its HTTP method, URL path, request parameters, and response format. All endpoints return JSON responses with a consistent schema containing 'status', 'message', and 'data' fields.")
blank(doc)

tbl(doc,
    ["#", "Method", "Endpoint", "Request Body / Params", "Response"],
    [
        ["1",  "POST", "/auth/send-otp",        "phone (string)",                          "200: {otp_sent: true} / 400: error"],
        ["2",  "POST", "/auth/verify-otp",       "phone, otp (string)",                     "200: {valid: true, token} / 401: invalid"],
        ["3",  "POST", "/auth/register",         "phone, name, role, password",             "201: {user_id, role} / 400: error"],
        ["4",  "POST", "/auth/login",            "phone, password",                         "200: {token, role, name} / 401: error"],
        ["5",  "GET",  "/auth/profile",          "Authorization header",                    "200: {name, phone, role, created_at}"],
        ["6",  "POST", "/disease/detect",        "image file (multipart/form-data)",        "200: {disease, confidence, heatmap_b64}"],
        ["7",  "POST", "/soil/analyse",          "nitrogen, phosphorus, potassium, ph, moisture", "200: {crop, confidence, health_score, advice[]}"],
        ["8",  "GET",  "/weather",               "city (query param)",                      "200: {current:{}, forecast:[]} / 404: city not found"],
        ["9",  "POST", "/chatbot",               "query (string)",                          "200: {response (string)} / 500: API error"],
        ["10", "GET",  "/marketplace/listings",  "status (optional query param)",           "200: [{listing_id, crop, qty, price, status, farmer}]"],
        ["11", "POST", "/marketplace/listings",  "crop_name, quantity_kg, price_per_kg",   "201: {listing_id} / 400: error"],
        ["12", "GET",  "/marketplace/listings/{id}", "listing_id (path param)",             "200: {listing detail + farmer info}"],
        ["13", "PUT",  "/marketplace/listings/{id}", "status (available/cancelled)",        "200: {updated} / 404: not found"],
        ["14", "POST", "/marketplace/buy",       "listing_id, quantity_kg",                "200: {purchase_id, total_price} / 400: error"],
        ["15", "GET",  "/marketplace/purchases", "Authorization header",                    "200: [{purchase detail}] — buyer's orders"],
        ["16", "GET",  "/messages/{listing_id}", "listing_id (path param)",                 "200: [{sender, content, timestamp}]"],
        ["17", "POST", "/messages/{listing_id}", "content (string)",                        "201: {message_id} / 400: error"],
        ["18", "GET",  "/finance",               "Authorization header",                    "200: [{type, amount, description, date}]"],
        ["19", "POST", "/finance",               "type (income/expense), amount, description", "201: {record_id}"],
        ["20", "GET",  "/calendar",              "Authorization header",                    "200: [{task, crop, due_date, completed}]"],
        ["21", "POST", "/calendar/generate",     "crop_name, sowing_date",                 "201: {tasks_created (int)}"],
        ["22", "PUT",  "/calendar/{task_id}",    "completed (boolean)",                    "200: {updated} / 404: not found"],
        ["23", "GET",  "/schemes",               "state (optional query param)",            "200: [{scheme_name, description, eligibility, url}]"],
    ],
    "Table B.1 Complete KrishiMitra REST API Endpoint Reference",
    [Cm(0.8), Cm(1.2), Cm(4.5), Cm(5.0), Cm(5.5)]
)

br(doc)

# ════════════════════════════════════════════
# APPENDIX C — Dataset Summary
# ════════════════════════════════════════════
sec(doc, "Appendix C: Training Dataset Summaries")
blank(doc)

sub(doc, "C.1 PlantVillage Dataset — CNN Disease Detection")
tbl(doc,
    ["Attribute", "Details"],
    [
        ["Dataset Name", "PlantVillage Disease Classification Dataset"],
        ["Source", "Hughes, D.P. and Salathe, M. (2015). arXiv:1511.08060"],
        ["Access", "TensorFlow Datasets: tfds.load('plant_village') — Open Access"],
        ["Total Images", "54,306 labelled leaf images"],
        ["Total Classes", "38 (26 disease classes + 12 healthy classes across 14 crop species)"],
        ["Classes Used in KrishiMitra", "3 classes: Healthy, Leaf_Spot, Rust (subset for prototype)"],
        ["Images Used (3 classes)", "Approx. 10,900 images"],
        ["Train/Validation Split", "80% training (8,720 images) / 20% validation (2,180 images)"],
        ["Image Format", "JPEG, RGB, variable resolution (resized to 224×224 for training)"],
        ["Data Augmentation Applied", "Random horizontal flip, rotation (±15°), zoom (±10%), brightness (±10%)"],
        ["Class Balance", "Approximately balanced across 3 selected classes"],
        ["Licence", "Open Access — free for academic and research use"],
        ["URL", "https://www.tensorflow.org/datasets/catalog/plant_village"],
    ],
    "Table C.1 PlantVillage Dataset Summary",
    [Cm(5), Cm(12)]
)

sub(doc, "C.2 Soil Crop Recommendation Dataset — KNN Model")
tbl(doc,
    ["Attribute", "Details"],
    [
        ["Dataset Name", "Crop Recommendation Dataset"],
        ["Source", "Kaggle — Agricultural dataset (Atharva Ingle, 2020)"],
        ["Access", "https://www.kaggle.com/datasets/atharvaingle/crop-recommendation-dataset"],
        ["Total Samples", "2,200 rows"],
        ["Total Crop Classes", "22 (Rice, Maize, Chickpea, Kidney Beans, Pigeon Peas, Moth Beans, Mung Bean, Blackgram, Lentil, Pomegranate, Banana, Mango, Grapes, Watermelon, Muskmelon, Apple, Orange, Papaya, Coconut, Cotton, Jute, Coffee)"],
        ["Features (Input)", "5 numerical features: N (Nitrogen kg/ha), P (Phosphorus kg/ha), K (Potassium kg/ha), pH (3.0–10.0), Humidity/Moisture (%)"],
        ["Label (Output)", "Crop name (categorical string)"],
        ["Samples per Class", "100 samples per crop class (balanced dataset)"],
        ["Train/Test Split", "80% training (1,760 samples) / 20% testing (440 samples)"],
        ["Normalisation", "Min-max normalisation applied to all 5 features before KNN training"],
        ["KNN Parameter", "k=7 (7 nearest neighbours — selected by cross-validation)"],
        ["File Format", "CSV (Crop_recommendation.csv)"],
        ["Licence", "CC0: Public Domain"],
    ],
    "Table C.2 Soil Crop Recommendation Dataset Summary",
    [Cm(5), Cm(12)]
)

br(doc)

# ════════════════════════════════════════════
# APPENDIX D — Project Screenshots / Scenarios
# ════════════════════════════════════════════
sec(doc, "Appendix D: Project Screenshots — Key Scenarios")
body(doc, "The following subsections describe the key application screens of KrishiMitra, demonstrating the system functioning across different user scenarios. Screenshots are taken from the production APK installed on a Realme C55 Android device (Android 13).")
blank(doc)

scenarios = [
    ("D.1", "Scenario 1 — User Registration and OTP Verification",
     "The user opens KrishiMitra for the first time. The Login screen is displayed with a phone number input field. The user enters their 10-digit phone number and taps 'Send OTP'. A 6-digit OTP is delivered via SMS through the Fast2SMS gateway within 2 seconds. The user enters the OTP in the verification field and taps 'Verify'. Upon successful verification, the Role Selection screen is displayed, where the user selects either 'Farmer' or 'Buyer'. The selection is saved to SharedPreferences for subsequent logins."),
    ("D.2", "Scenario 2 — Crop Disease Detection with Grad-CAM",
     "The Farmer navigates to the Disease Detection screen from the Farmer Dashboard. The screen displays a camera icon and an 'Upload Image' button. The farmer taps the camera icon, which opens the device camera via the image_picker plugin. The farmer photographs a diseased tomato leaf showing rust-coloured pustules. The image is displayed on screen with an 'Analyse' button. Upon tapping, a CircularProgressIndicator is shown while the image is sent to the /disease/detect Flask endpoint. The API returns within 2.3 seconds with the result: 'Rust — 91.2% confidence'. The Grad-CAM heatmap overlay is displayed below, highlighting the rust pustule regions in red/yellow, clearly showing which areas drove the diagnosis."),
    ("D.3", "Scenario 3 — Soil Analysis and Crop Recommendation",
     "The Farmer navigates to the Soil Analysis screen. A form with five input fields is displayed: Nitrogen (kg/ha), Phosphorus (kg/ha), Potassium (kg/ha), Soil pH, and Moisture (%). The farmer enters measured values: N=90, P=42, K=43, pH=6.5, Moisture=60%. Upon tapping 'Analyse Soil', the values are sent to the /soil/analyse endpoint. The result screen displays: Recommended Crop: Rice (83.4% confidence), Soil Health Score: 74/100 (Good), and Fertiliser Advice: 'Apply MOP 40kg/ha to boost Potassium'. The health score is displayed as a colour-coded progress bar (green for >70%)."),
    ("D.4", "Scenario 4 — P2P Marketplace — Farmer Lists Crop",
     "The Farmer navigates to the Marketplace screen and taps the '+' floating action button to create a new listing. A form is displayed with fields for Crop Name, Quantity (kg), and Price per kg (INR). The farmer enters: Rice, 500 kg, ₹28/kg. Upon submission, the listing is created via the /marketplace/listings POST endpoint and appears immediately in the marketplace feed with status 'Available'."),
    ("D.5", "Scenario 5 — P2P Marketplace — Buyer Purchases Crop",
     "The Buyer navigates to the Marketplace screen, which shows all active listings. The buyer taps a Rice listing (500 kg at ₹28/kg). The Listing Detail screen shows full information including farmer name, location, and a chat icon. The buyer taps 'Buy Now', enters quantity 100 kg, and confirms the purchase. The total price (₹2,800) is calculated and displayed. Upon confirmation, the listing status changes to 'Purchased' and a chat thread opens for negotiation between the buyer and farmer."),
    ("D.6", "Scenario 6 — AI Chatbot Consultation",
     "The Farmer navigates to the AI Chatbot screen. A chat-style interface is displayed with a text input at the bottom. The farmer types: 'What fertiliser should I use for paddy crop showing yellow leaves?' The query is sent to the /chatbot endpoint which proxies it to the Groq LLaMA 3.3-70B API. Within 3.1 seconds, a detailed agronomic response is returned, recommending Urea application at 60 kg/ha for Nitrogen deficiency (indicated by yellowing leaves) and advising soil testing. The response is displayed in a chat bubble with a green KrishiMitra logo avatar."),
    ("D.7", "Scenario 7 — Weather Advisory",
     "The Farmer navigates to the Weather screen. The app detects the device location (or uses a saved city) and queries the OpenWeatherMap API. The current conditions are displayed: 28°C, 72% humidity, Light Rain, Wind 12 km/h. Below the current conditions, a 5-day forecast strip is shown with weather icons. The agriculture advisory section below displays: 'Avoid pesticide spraying today — rain forecast may wash off chemicals. Good conditions for irrigation — no need for supplemental watering in next 48 hours.'"),
    ("D.8", "Scenario 8 — Finance Tracker",
     "The Farmer navigates to the Finance Tracker screen. A summary card at the top shows: Total Income this month: ₹14,500 | Total Expenses: ₹3,200 | Net: ₹11,300. Below, a list of recent transactions is shown with colour coding (green for income, red for expense). The farmer taps '+' to add a new expense: 'Urea fertiliser purchase — ₹850'. The new entry appears in the list and the summary card updates automatically."),
]

for code, scenario_title, description in scenarios:
    sub(doc, f"{code} {scenario_title}")
    body(doc, description)
    cap(doc, f"[Screenshot: {scenario_title}]")
    blank(doc)

br(doc)

# ════════════════════════════════════════════
# APPENDIX E — Installation & Configuration Guide
# ════════════════════════════════════════════
sec(doc, "Appendix E: Installation and Configuration Guide")
blank(doc)

sub(doc, "E.1 Backend Setup (Flask + Python)")
for step in [
    "Step 1: Install Python 3.10 or higher from https://www.python.org/downloads/",
    "Step 2: Navigate to the krishimitra_backend/ directory",
    "Step 3: Install dependencies: pip install -r requirements.txt",
    "Step 4: Set environment variables: FAST2SMS_API_KEY, GROQ_API_KEY, OPENWEATHER_API_KEY",
    "Step 5: Run the server: python app.py",
    "Step 6: The Flask API will start at http://0.0.0.0:5000",
    "Step 7: Verify by visiting http://localhost:5000/health in a browser — should return {status: ok}",
]:
    bullet(doc, step)
blank(doc)

sub(doc, "E.2 Frontend Setup (Flutter)")
for step in [
    "Step 1: Install Flutter SDK 3.19+ from https://flutter.dev/docs/get-started/install",
    "Step 2: Install Android Studio and Android SDK (API Level 21+)",
    "Step 3: Navigate to the krishimitra_app/ directory",
    "Step 4: Update the backend IP address in lib/services/api_service.dart (line: baseUrl = 'http://YOUR_IP:5000')",
    "Step 5: Run: flutter pub get  (installs all Dart package dependencies)",
    "Step 6: Connect Android device via USB with Developer Mode and USB Debugging enabled",
    "Step 7: Run: flutter run  (for debug build) or flutter build apk --release (for production APK)",
    "Step 8: The APK file will be generated at build/app/outputs/flutter-apk/app-release.apk",
]:
    bullet(doc, step)
blank(doc)

sub(doc, "E.3 Key Configuration Files")
tbl(doc,
    ["File", "Location", "Purpose"],
    [
        ["requirements.txt",   "krishimitra_backend/", "Python dependency list — Flask, TensorFlow, SQLAlchemy, OpenCV, etc."],
        ["app.py",             "krishimitra_backend/", "Flask application factory — initialises app, registers blueprints, creates DB"],
        ["api_service.dart",   "krishimitra_app/lib/services/", "Flutter API client — configure baseUrl to match Flask server IP:port"],
        ["pubspec.yaml",       "krishimitra_app/", "Flutter dependency manifest — all Dart/Flutter packages"],
        ["build.gradle.kts",   "krishimitra_app/android/app/", "Android build config — minSdkVersion, targetSdkVersion, app ID"],
        ["disease_model.h5",   "krishimitra_backend/models/", "Pre-trained ResNet50 CNN model file — required for disease detection"],
        ["soil_data.json",     "krishimitra_backend/data/", "KNN training data — 2200 soil samples with crop labels"],
        ["krishimitra.db",     "krishimitra_backend/instance/", "SQLite database file — auto-created on first run"],
    ],
    "Table E.1 Key Configuration Files",
    [Cm(4), Cm(4), Cm(9)]
)

doc.save("KrishiMitra_Appendix.docx")
print("Saved: KrishiMitra_Appendix.docx")
print("Sections: A (Library Specs), B (API Reference), C (Datasets), D (Screenshots), E (Installation)")
