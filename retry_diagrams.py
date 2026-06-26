import urllib.request
import base64
import time
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def download_mermaid_ink(mermaid_str, filename):
    encoded = base64.b64encode(mermaid_str.encode('utf-8')).decode('ascii')
    url = f'https://mermaid.ink/img/{encoded}'
    
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'})
    
    for attempt in range(5):
        print(f'Downloading {filename} (Attempt {attempt+1})...')
        try:
            with urllib.request.urlopen(req, timeout=30) as response, open(filename, 'wb') as out_file:
                out_file.write(response.read())
            print(f'Success: {filename}')
            return True
        except Exception as e:
            print(f'Failed {filename}: {e}')
            time.sleep(10)
    return False

er_code = """
erDiagram
    USER ||--o{ LISTING : "creates"
    USER ||--o{ PURCHASE : "makes"
    USER ||--o{ MESSAGE : "sends"
    USER ||--o{ TRANSACTION : "has"
    USER ||--o{ CALENDARTASK : "manages"
    USER ||--o{ ECOMMERCEORDER : "places"
    LISTING ||--o{ PURCHASE : "contains"
    PURCHASE ||--o{ MESSAGE : "tracks"
"""

flowchart_code = """
graph TD
    A([Start]) --> B[Capture Leaf Image]
    B --> C[Compress & Encode Base64]
    C --> D[Transmit via HTTPS POST]
    D --> E{Backend Flask API}
    E --> F[Decode Image Tensor]
    F --> G[ResNet50 Inference]
    G --> H[Extract Feature Maps]
    H --> I[Apply Grad-CAM]
    I --> J[Overlay JET Colormap]
    J --> K[Return JSON + Image URL]
    K --> L[Render Heatmap]
    L --> M([End])
"""

diagrams = [
    (er_code, "real_er.png", "Entity-Relationship (ER) Diagram", "5.4 Database Design and Data Dictionary", 3),
    (flowchart_code, "real_flow.png", "Flowchart: Disease Detection Pipeline", "6.2.1 Deep Learning Vision Module", 4)
]

for code, filename, caption, target_text, fig_num in diagrams:
    if not os.path.exists(filename):
        download_mermaid_ink(code, filename)

# Re-open the partially updated doc
doc_path = "KrishiMitra_Final_Report_Massive_Correct_Diagrams.docx"
print(f"Opening {doc_path}...")
try:
    doc = Document(doc_path)
except Exception as e:
    print(f"Error opening doc: {e}")
    exit(1)

for code, filename, caption, target_text, fig_num in diagrams:
    if not os.path.exists(filename):
        continue
        
    # Check if already inserted
    already_inserted = False
    for p in doc.paragraphs:
        if f"Figure {fig_num}: {caption}" in p.text:
            already_inserted = True
            break
    if already_inserted:
        print(f"Already inserted: {filename}")
        continue
        
    for i, p in enumerate(doc.paragraphs):
        if target_text in p.text:
            print(f"Found target: {target_text}")
            insert_idx = min(i + 2, len(doc.paragraphs) - 1)
            target_p = doc.paragraphs[insert_idx]
            
            new_p = target_p.insert_paragraph_before()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_p.add_run()
            run.add_picture(filename, width=Inches(5.0))
            
            cap_p = target_p.insert_paragraph_before()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_r = cap_p.add_run(f"Figure {fig_num}: {caption}")
            cap_r.font.size = Pt(12)
            cap_r.font.name = "Times New Roman"
            cap_r.bold = True
            
            target_p.insert_paragraph_before("")
            break

out_name = "KrishiMitra_Final_Report_Massive_Correct_Diagrams.docx"
doc.save(out_name)
print(f"Successfully updated {out_name} with missing diagrams!")
